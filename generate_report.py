"""
AquaSense Report Generator  —  python generate_report.py
Produces AquaSense_Report.docx with professional academic design.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ═══════════════════════════════════════════════════════════
#  COLOUR PALETTE
# ═══════════════════════════════════════════════════════════
C_GREEN_DARK   = "065F46"
C_GREEN_MID    = "10B981"
C_GREEN_LIGHT  = "ECFDF5"
C_GREEN_PALE   = "F0FDF4"
C_BLUE_DARK    = "1E3A5F"
C_BLUE_MID     = "1D4ED8"
C_BLUE_LIGHT   = "EFF6FF"
C_INDIGO       = "4338CA"
C_INDIGO_LIGHT = "EEF2FF"
C_GRAY_DARK    = "1F2937"
C_GRAY_MID     = "374151"
C_GRAY_LIGHT   = "6B7280"
C_GRAY_PALE    = "F8FAFC"
C_WHITE        = "FFFFFF"

# ═══════════════════════════════════════════════════════════
#  LOW-LEVEL XML HELPERS
# ═══════════════════════════════════════════════════════════

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove existing shd
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color, sz="12", sides=("top","left","bottom","right")):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for side in sides:
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    sz)
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def clear_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        b.set(qn("w:sz"),  "0")
        b.set(qn("w:color"), "auto")
        tcBorders.append(b)
    tcPr.append(tcBorders)

def set_run_color(run, hex_color):
    run.font.color.rgb = RGBColor(
        int(hex_color[0:2], 16),
        int(hex_color[2:4], 16),
        int(hex_color[4:6], 16),
    )

def para_border(para, side, color, sz="18"):
    pPr  = para._p.get_or_add_pPr()
    pBdr = pPr.find(qn("w:pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"),   "single")
    el.set(qn("w:sz"),    sz)
    el.set(qn("w:space"), "4")
    el.set(qn("w:color"), color)
    pBdr.append(el)

def set_para_shading(para, fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill)
    pPr.append(shd)

def set_cell_vertical_align(cell, align="center"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)

def add_page_num_to_footer(footer):
    """Insert 'Page X of Y' centred in a footer paragraph."""
    para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(4)

    def _fld(instr):
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), instr)
        return fld
    def _instr(text):
        ri = OxmlElement("w:instrText")
        ri.set(qn("xml:space"), "preserve")
        ri.text = text
        return ri

    r1 = para.add_run("Page ")
    r1.font.size = Pt(9)
    set_run_color(r1, C_GRAY_LIGHT)

    rn = OxmlElement("w:r")
    rn.append(_fld("begin"))
    rn.append(_instr(" PAGE "))
    rn.append(_fld("separate"))
    rn.append(_fld("end"))
    para._p.append(rn)

    r2 = para.add_run(" of ")
    r2.font.size = Pt(9)
    set_run_color(r2, C_GRAY_LIGHT)

    rp = OxmlElement("w:r")
    rp.append(_fld("begin"))
    rp.append(_instr(" NUMPAGES "))
    rp.append(_fld("separate"))
    rp.append(_fld("end"))
    para._p.append(rp)

def add_header_text(header, text):
    para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    para.paragraph_format.space_after = Pt(4)
    para_border(para, "bottom", C_GREEN_MID, sz="6")
    run = para.add_run(text)
    run.font.size = Pt(8.5)
    run.italic = True
    set_run_color(run, C_GRAY_LIGHT)

# ═══════════════════════════════════════════════════════════
#  HIGH-LEVEL CONTENT HELPERS
# ═══════════════════════════════════════════════════════════

def styled_run(para, text, size=11, bold=False, italic=False, color=C_GRAY_MID):
    r = para.add_run(text)
    r.bold        = bold
    r.italic      = italic
    r.font.size   = Pt(size)
    r.font.name   = "Calibri"
    set_run_color(r, color)
    return r

def add_para(doc, text="", bold=False, italic=False, size=11,
             color=C_GRAY_MID, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        styled_run(p, text, size=size, bold=bold, italic=italic, color=color)
    return p

def section_heading(doc, number, title, color=C_GREEN_DARK):
    """
    Styled section heading: coloured left-bar accent + number pill.
    Uses the built-in Heading 1 style so TOC can pick it up.
    """
    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    # left border accent
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "single")
    left.set(qn("w:sz"),    "28")
    left.set(qn("w:space"), "10")
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)
    # text
    r = p.add_run(f"{number}  {title}")
    r.bold = True
    r.font.size = Pt(14)
    r.font.name = "Calibri"
    set_run_color(r, color)
    return p

def sub_heading(doc, text, color=C_BLUE_MID):
    """Heading 2 equivalent with green underline."""
    p = doc.add_paragraph(style="Heading 2")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = "Calibri"
    set_run_color(r, color)
    return p

def sub_sub_heading(doc, text, color=C_GRAY_DARK):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    set_run_color(r, color)
    return p

def add_bullet(doc, text, bold_prefix="", color=C_GRAY_MID, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.bold = True
        rb.font.size = Pt(10.5)
        rb.font.name = "Calibri"
        set_run_color(rb, C_GRAY_DARK)
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = "Calibri"
    set_run_color(r, color)
    return p

def info_box(doc, title, body, bg=C_BLUE_LIGHT, border=C_BLUE_MID):
    """Highlighted callout box."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_borders(cell, border, sz="12")
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_before = Pt(6)
    p1.paragraph_format.space_after  = Pt(2)
    p1.paragraph_format.left_indent  = Cm(0.3)
    rt = p1.add_run(title + "  ")
    rt.bold = True
    rt.font.size = Pt(10)
    rt.font.name = "Calibri"
    set_run_color(rt, border)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after  = Pt(6)
    p2.paragraph_format.left_indent  = Cm(0.3)
    rb = p2.add_run(body)
    rb.font.size = Pt(10)
    rb.italic    = True
    rb.font.name = "Calibri"
    set_run_color(rb, C_GRAY_MID)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def screenshot_box(doc, label, hint=""):
    tbl  = doc.add_table(rows=1, cols=1)
    tbl.style  = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, C_INDIGO_LIGHT)
    set_cell_borders(cell, C_INDIGO, sz="18")

    cell.add_paragraph().paragraph_format.space_after = Pt(4)

    p_icon = cell.add_paragraph()
    p_icon.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ri = p_icon.add_run("[  SCREENSHOT  ]")
    ri.bold = True
    ri.font.size = Pt(9)
    ri.font.name = "Calibri"
    set_run_color(ri, C_INDIGO)

    p_lbl = cell.add_paragraph()
    p_lbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rl = p_lbl.add_run(label)
    rl.bold = True
    rl.font.size = Pt(12)
    rl.font.name = "Calibri"
    set_run_color(rl, "3730A3")

    if hint:
        p_h = cell.add_paragraph()
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rh = p_h.add_run(hint)
        rh.italic    = True
        rh.font.size = Pt(9)
        rh.font.name = "Calibri"
        set_run_color(rh, C_GRAY_LIGHT)

    cell.add_paragraph().paragraph_format.space_after = Pt(4)
    doc.add_paragraph().paragraph_format.space_after  = Pt(6)
    return tbl

def data_table(doc, rows_data, headers, hdr_bg=C_BLUE_DARK):
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_cell_bg(c, hdr_bg)
        set_cell_vertical_align(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.name = "Calibri"
        set_run_color(r, C_WHITE)

    for ri, row_vals in enumerate(rows_data):
        bg = C_GRAY_PALE if ri % 2 == 0 else C_WHITE
        for ci, val in enumerate(row_vals):
            c = tbl.rows[ri + 1].cells[ci]
            set_cell_bg(c, bg)
            set_cell_vertical_align(c)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after  = Pt(4)
            r = p.add_run(str(val))
            r.font.size = Pt(10)
            r.font.name = "Calibri"
            set_run_color(r, C_GRAY_DARK)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    para_border(p, "bottom", C_GREEN_MID, sz="6")
    return p

# ═══════════════════════════════════════════════════════════
#  BUILD DOCUMENT
# ═══════════════════════════════════════════════════════════

doc = Document()

# ── Page setup ───────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.2)
    sec.bottom_margin = Cm(2.2)
    sec.left_margin   = Cm(2.8)
    sec.right_margin  = Cm(2.2)

# ── Default font ─────────────────────────────────────────
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(11)

# ════════════════════════════════════════════════════════════════════════════
#  SECTION 0 — COVER PAGE  (no header/footer on this section)
# ════════════════════════════════════════════════════════════════════════════

cover_section = doc.sections[0]
cover_section.different_first_page_header_footer = True

# ── Green top banner ─────────────────────────────────────
top_bar = doc.add_table(rows=1, cols=1)
top_bar.alignment = WD_TABLE_ALIGNMENT.CENTER
tc_top = top_bar.rows[0].cells[0]
set_cell_bg(tc_top, C_GREEN_DARK)
clear_cell_borders(tc_top)

p_univ = tc_top.paragraphs[0]
p_univ.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_univ.paragraph_format.space_before = Pt(16)
p_univ.paragraph_format.space_after  = Pt(4)
styled_run(p_univ, "UNIVERSITY OF RWANDA",
           size=22, bold=True, color=C_WHITE)

p_coll = tc_top.add_paragraph()
p_coll.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_coll.paragraph_format.space_after = Pt(16)
styled_run(p_coll, "College of Science and Technology",
           size=12, color="D1FAE5")

# ── Institution lines ─────────────────────────────────────
doc.add_paragraph().paragraph_format.space_after = Pt(6)

def cv(t, sz=11, b=False, col=C_GRAY_MID, sa=3):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(sa)
    styled_run(p, t, size=sz, bold=b, color=col)
    return p

cv("School of Information and Communication Technology",
   sz=12, b=True, col=C_BLUE_MID, sa=3)
cv("Department of Information Technology",
   sz=11, col=C_GRAY_MID, sa=12)

# divider line
div_top = doc.add_paragraph()
div_top.paragraph_format.space_after = Pt(10)
para_border(div_top, "bottom", C_GREEN_DARK, sz="24")

# ── Report-type badge ─────────────────────────────────────
badge = doc.add_table(rows=1, cols=1)
badge.alignment = WD_TABLE_ALIGNMENT.CENTER
bc = badge.rows[0].cells[0]
set_cell_bg(bc, C_BLUE_DARK)
clear_cell_borders(bc)
bp = bc.paragraphs[0]
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
bp.paragraph_format.space_before = Pt(7)
bp.paragraph_format.space_after  = Pt(7)
styled_run(bp, "   PROJECT REPORT   ", size=13, bold=True, color=C_WHITE)

doc.add_paragraph().paragraph_format.space_after = Pt(6)
cv("Course: EIT4151 — Ubiquitous and Pervasive Computing (IoT)",
   sz=11, col=C_GRAY_MID, sa=14)

# ── Title block ───────────────────────────────────────────
title_tbl = doc.add_table(rows=1, cols=1)
title_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tc_title = title_tbl.rows[0].cells[0]
set_cell_bg(tc_title, C_GREEN_LIGHT)
set_cell_borders(tc_title, C_GREEN_DARK, sz="24")

p_aq = tc_title.paragraphs[0]
p_aq.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_aq.paragraph_format.space_before = Pt(16)
p_aq.paragraph_format.space_after  = Pt(4)
styled_run(p_aq, "AquaSense", size=40, bold=True, color=C_GREEN_DARK)

p_sub = tc_title.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(16)
styled_run(p_sub, "Smart Irrigation System Using IoT and Machine Learning",
           size=15, bold=True, color=C_BLUE_MID)

doc.add_paragraph().paragraph_format.space_after = Pt(10)
cv("Submitted in partial fulfilment of the requirements for the IoT module",
   sz=10, col=C_GRAY_LIGHT, sa=14)

# ── Members table ─────────────────────────────────────────
members = [
    ("1",  "SIBONIYO Emmanuel",            "222006224"),
    ("2",  "NZARAMYIMANA Jerome",          "222008510"),
    ("3",  "BYUKUSENGE Immaculee",         "222006273"),
    ("4",  "MUKAGASIRABO Beatrice",        "222004462"),
    ("5",  "UWIRINGIYIMANA Marie Claire",  "222004637"),
    ("6",  "TUYISHIME Ephron",             "222005571"),
    ("7",  "DUSHIMIMANA Fabrice",          "222017059"),
    ("8",  "MUGABO KAZINA Jules",          "222002936"),
    ("9",  "ISHIMWE Jean Marie Vianney",   "222019273"),
    ("10", "MUGISHA Edson",                "222018513"),
    ("11", "IRADUKUNDA Olivier",           "222005508"),
]

mem_tbl = doc.add_table(rows=1 + len(members), cols=3)
mem_tbl.style = "Table Grid"
mem_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# column widths
col_w = [Cm(1.4), Cm(9.2), Cm(4.4)]
for row in mem_tbl.rows:
    for i, w in enumerate(col_w):
        row.cells[i].width = w

# header row
for i, (txt, bg) in enumerate(
        zip(["No.", "Full Name", "Registration Number"],
            [C_GREEN_DARK, C_GREEN_DARK, C_GREEN_DARK])):
    c = mem_tbl.rows[0].cells[i]
    set_cell_bg(c, bg)
    set_cell_vertical_align(c)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(11)
    r.font.name = "Calibri"
    set_run_color(r, C_WHITE)

# data rows
for ri, (num, name, reg) in enumerate(members):
    bg = C_GREEN_PALE if ri % 2 == 0 else C_WHITE
    row = mem_tbl.rows[ri + 1]

    # No.
    c0 = row.cells[0]
    set_cell_bg(c0, bg)
    set_cell_vertical_align(c0)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(5)
    p0.paragraph_format.space_after  = Pt(5)
    r0 = p0.add_run(num)
    r0.font.size = Pt(11)
    r0.font.name = "Calibri"
    set_run_color(r0, C_GRAY_MID)

    # Name
    c1 = row.cells[1]
    set_cell_bg(c1, bg)
    set_cell_vertical_align(c1)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_before = Pt(5)
    p1.paragraph_format.space_after  = Pt(5)
    p1.paragraph_format.left_indent  = Cm(0.4)
    r1 = p1.add_run(name)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.name = "Calibri"
    set_run_color(r1, C_GRAY_DARK)

    # Reg
    c2 = row.cells[2]
    set_cell_bg(c2, bg)
    set_cell_vertical_align(c2)
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(5)
    p2.paragraph_format.space_after  = Pt(5)
    r2 = p2.add_run(reg)
    r2.font.size = Pt(11)
    r2.font.name = "Calibri"
    set_run_color(r2, C_GREEN_DARK)

doc.add_paragraph().paragraph_format.space_after = Pt(12)

# ── Bottom divider ────────────────────────────────────────
div_bot = doc.add_paragraph()
div_bot.paragraph_format.space_after = Pt(8)
para_border(div_bot, "top", C_GREEN_DARK, sz="24")

# ── Footer meta (3-col borderless) ───────────────────────
foot_meta = doc.add_table(rows=1, cols=3)
foot_meta.alignment = WD_TABLE_ALIGNMENT.CENTER
foot_items = [
    ("Supervisor:", "IoT Course Lecturer"),
    ("Academic Year:", "2025 / 2026"),
    ("Date:", "May 2026"),
]
for i, (lbl, val) in enumerate(foot_items):
    c = foot_meta.rows[0].cells[i]
    clear_cell_borders(c)
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    rl = p.add_run(lbl + "  ")
    rl.bold = True
    rl.font.size = Pt(10)
    rl.font.name = "Calibri"
    set_run_color(rl, C_GRAY_DARK)
    rv = p.add_run(val)
    rv.font.size = Pt(10)
    rv.font.name = "Calibri"
    set_run_color(rv, C_GRAY_MID)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ── Green bottom banner ───────────────────────────────────
bot_bar = doc.add_table(rows=1, cols=1)
bot_bar.alignment = WD_TABLE_ALIGNMENT.CENTER
tc_bot = bot_bar.rows[0].cells[0]
set_cell_bg(tc_bot, C_GREEN_DARK)
clear_cell_borders(tc_bot)
p_bot = tc_bot.paragraphs[0]
p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_bot.paragraph_format.space_before = Pt(10)
p_bot.paragraph_format.space_after  = Pt(10)
styled_run(p_bot,
    "University of Rwanda  |  Kigali, Rwanda  |  Academic Year 2025-2026",
    size=10, color="D1FAE5")

# ════════════════════════════════════════════════════════════════════════════
#  NEW SECTION — header + footer start here (all pages after cover)
# ════════════════════════════════════════════════════════════════════════════

doc.add_page_break()

# Add a new section so we can have header/footer only from page 2 onward
from docx.oxml import OxmlElement as OE
new_sec_props = OxmlElement("w:sectPr")
new_sec_props.set(qn("w:type"), "nextPage")
doc.paragraphs[-1]._p.append(new_sec_props)

# Configure the second (main) section header & footer
main_section = doc.sections[-1] if len(doc.sections) > 1 else doc.sections[0]
main_section.header_distance = Cm(1.2)
main_section.footer_distance = Cm(1.2)

# Header
hdr = main_section.header
hdr.is_linked_to_previous = False
add_header_text(hdr, "AquaSense  |  Smart Irrigation IoT System  |  University of Rwanda")

# Footer
ftr = main_section.footer
ftr.is_linked_to_previous = False
add_page_num_to_footer(ftr)

# ════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS PAGE
# ════════════════════════════════════════════════════════════════════════════

p_toc_title = doc.add_paragraph()
p_toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_toc_title.paragraph_format.space_after = Pt(4)
set_para_shading(p_toc_title, C_GREEN_DARK)
rt = p_toc_title.add_run("  TABLE OF CONTENTS  ")
rt.bold = True
rt.font.size = Pt(16)
rt.font.name = "Calibri"
set_run_color(rt, C_WHITE)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

toc_entries = [
    ("Executive Summary",                        ""),
    ("1.  System Overview and Architecture",     ""),
    ("     1.1  Architecture Layers",            ""),
    ("2.  Hardware Components",                  ""),
    ("3.  Sensor Calibration and Signal Conversion", ""),
    ("     3.1  Soil Moisture (ADC)",            ""),
    ("     3.2  Temperature & Humidity (DHT11)", ""),
    ("     3.3  Reservoir Level (HC-SR04)",      ""),
    ("     3.4  NPK Sensor (RS-485 Modbus)",     ""),
    ("4.  Data Collection and Database",         ""),
    ("     4.1  Dataset Statistics",             ""),
    ("     4.2  Descriptive Statistics",         ""),
    ("5.  Dashboard — Main Page (/dashboard)",   ""),
    ("     5.1  Stats Bar",                      ""),
    ("     5.2  Live Sensor Overview",           ""),
    ("     5.3  Multi-Sensor Time-Series Chart", ""),
    ("     5.4  Conversion Formulas Panel",      ""),
    ("     5.5  Data Classification Panel",      ""),
    ("     5.6  AI Summary Card",                ""),
    ("     5.7  Node Cards + Pump Control",      ""),
    ("6.  Node Detail Page (/nodes/[id])",       ""),
    ("7.  AI Intelligence Page (/ml)",           ""),
    ("     7.1  Stats Bar & Hero",               ""),
    ("     7.2  Current Decision (Live)",        ""),
    ("     7.3  Live Arduino Command",           ""),
    ("     7.4  Model Performance Metrics",      ""),
    ("     7.5  Feature Importance & Decisions", ""),
    ("     7.6  Decision History Table",         ""),
    ("     7.7  Python Analysis Charts",         ""),
    ("     7.8  How the AI Model Works",         ""),
    ("8.  Python Analysis Pipeline",             ""),
    ("     8.1  Pipeline Modules",               ""),
    ("     8.2  Classification Results",         ""),
    ("     8.3  Regression Results",             ""),
    ("     8.4  Cross-Validation Comparison",    ""),
    ("     8.5  Generated Analysis Charts",      ""),
    ("9.  Results and Discussion",               ""),
    ("10. Conclusion",                           ""),
    ("11. References",                           ""),
]

for i, (entry, _) in enumerate(toc_entries):
    is_main = not entry.startswith("  ")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2 if not is_main else 5)
    if is_main:
        set_para_shading(p, C_GREEN_PALE)
        para_border(p, "left", C_GREEN_MID, sz="16")
        p.paragraph_format.left_indent = Cm(0.3)
    else:
        p.paragraph_format.left_indent = Cm(1.5)
    r = p.add_run(entry)
    r.bold = is_main
    r.font.size = Pt(11 if is_main else 10)
    r.font.name = "Calibri"
    set_run_color(r, C_GRAY_DARK if is_main else C_GRAY_MID)
    # dots filler
    dots = p.add_run("  " + ("." * (55 if is_main else 50)))
    dots.font.size = Pt(10)
    dots.font.name = "Calibri"
    set_run_color(dots, "D1D5DB")

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
#  EXECUTIVE SUMMARY
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "", "Executive Summary")
add_para(doc,
    "AquaSense is a full-stack IoT smart-irrigation platform built for precision "
    "agriculture. It collects real-time data from seven sensor channels on an Arduino "
    "Uno node, persists every reading in a PostgreSQL database, provides a professional "
    "Next.js 14 dashboard for live monitoring, and runs a Python/FastAPI machine-learning "
    "service that decides in real time whether to irrigate, hold, or flag a low-water "
    "condition. By the time of submission the system had accumulated 104,370 sensor "
    "readings and issued tens of thousands of irrigation decisions, with a Random Forest "
    "model achieving perfect in-distribution accuracy and 94.99% mean cross-validated F1.",
    size=11, space_after=8)

info_box(doc,
    "Key Achievement:",
    "104,370 sensor readings stored  |  Random Forest 94.99% F1  "
    "|  11 Python analysis charts  |  Real-time pump control via ML",
    bg=C_GREEN_LIGHT, border=C_GREEN_DARK)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
#  1. SYSTEM OVERVIEW
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "1.", "System Overview and Architecture")
add_para(doc,
    "AquaSense follows a four-tier IoT architecture. Each tier is a distinct "
    "process communicating via well-defined interfaces.", size=11, space_after=8)

sub_heading(doc, "1.1  Architecture Layers")

tiers = [
    ("Arduino Uno (Edge Layer)",
     "Reads all sensors every 2 s. Emits a DATA: line over USB serial. "
     "Receives CMD:PUMP=1/0 commands and drives the relay. Falls back to "
     "a local rule if no ML command arrives within 10 s."),
    ("Python Serial Bridge (Gateway Layer)",
     "Converts raw ADC/Modbus bytes to engineering units with calibration "
     "offsets, then HTTP-POSTs each reading to the Next.js API. A JSONL "
     "offline queue retries failed posts."),
    ("Next.js 14 + PostgreSQL (Backend & Dashboard)",
     "App Router server components fetch data from PostgreSQL via Prisma ORM "
     "and render a live dashboard. Auto-refreshes every 5 s. REST routes at "
     "/api/* handle ingest, pump events, and ML decision writes."),
    ("Python FastAPI ML Service (AI Engine)",
     "Port 5001: POST /predict, POST /train, GET /status. Trains a Random "
     "Forest on all DB readings, issues IRRIGATE/HOLD/LOW_WATER, writes "
     "results directly to PostgreSQL, and auto-retrains every 50 new readings."),
]
for title, body in tiers:
    add_bullet(doc, body, bold_prefix=title + ": ", color=C_GRAY_MID)

doc.add_paragraph().paragraph_format.space_after = Pt(4)
info_box(doc,
    "Data Flow:",
    "Sensors -> Arduino (ADC / RS-485 / DHT11 / HC-SR04) -> "
    "Serial Bridge (calibration) -> Next.js API -> PostgreSQL -> "
    "ML Service (Random Forest) -> IrrigationDecision table -> "
    "Dashboard + Arduino relay command (feedback loop)",
    bg=C_BLUE_LIGHT, border=C_BLUE_MID)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
#  2. HARDWARE
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "2.", "Hardware Components")
data_table(doc, [
    ["Arduino Uno R3",           "Main MCU", "14 digital I/O, 6 analogue inputs, USB serial"],
    ["Capacitive Soil Moisture", "Analogue A0", "0-1023 ADC, lower = wetter"],
    ["DHT11",                    "Digital pin 7", "Temperature 0-50 degC, Humidity 20-90% RH"],
    ["HC-SR04 Ultrasonic",       "TRIG=9, ECHO=10", "2-400 cm; used for reservoir level"],
    ["RS-485 NPK Sensor",        "SoftwareSerial RX=2 TX=3", "Modbus RTU, N/P/K in mg/kg"],
    ["5V Relay Module",          "Digital pin 8", "Controls 12V pump; HIGH = pump ON"],
    ["LED (built-in pin 13)",    "Digital pin 13", "Visual pump status indicator"],
], ["Component", "Interface", "Notes"], hdr_bg=C_GREEN_DARK)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
#  3. CALIBRATION
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "3.", "Sensor Calibration and Signal Conversion")
add_para(doc,
    "Every sensor delivers a raw electrical signal that must be converted to a "
    "physically meaningful unit. Conversions are shown on the dashboard's "
    "'Conversion Formulas' panel and documented here.", size=11, space_after=8)

sub_heading(doc, "3.1  Soil Moisture — Capacitive ADC")
add_bullet(doc, "Raw signal: 0-1023 (10-bit ADC, pin A0)")
add_bullet(doc, "Dry calibration: raw_dry = 820   |   Wet calibration: raw_wet = 390")
add_bullet(doc, "Formula:  moisture_pct = (raw_dry - ADC) / (raw_dry - raw_wet) x 100")
add_bullet(doc, "Pump threshold: moisture < 78% triggers IRRIGATE  (raw threshold = 225)")
add_bullet(doc, "Output clamped to [0, 100] %")

sub_heading(doc, "3.2  Temperature & Humidity — DHT11")
add_bullet(doc, "DHT11 outputs calibrated digital values natively — no ADC conversion needed")
add_bullet(doc, "Bridge applies configurable offset: temp +0.0 degC, humidity +0.0% RH")
add_bullet(doc, "Valid range: temperature -40 to 80 degC; humidity 0-100% RH")

sub_heading(doc, "3.3  Reservoir Level — HC-SR04 Ultrasonic")
add_bullet(doc, "TANK_FULL_CM = 3.0 cm (sensor-to-water when full), TANK_EMPTY_CM = 12.0 cm")
add_bullet(doc, "Formula:  level_pct = (TANK_EMPTY_CM - distance) / (TANK_EMPTY_CM - TANK_FULL_CM) x 100")
add_bullet(doc, "LOW_WATER: distance > 11.0 cm (< 15% full) -> pump disabled")

sub_heading(doc, "3.4  NPK Sensor — RS-485 Modbus RTU")
add_bullet(doc, "SoftwareSerial at 4800 baud, Modbus function code 0x03")
add_bullet(doc, "Registers 0x0000-0x0002: N, P, K — 16-bit unsigned integer = mg/kg directly")
add_bullet(doc, "Valid range 0-1999 mg/kg per channel; clamped by bridge")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
#  4. DATA
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "4.", "Data Collection and Database")

sub_heading(doc, "4.1  Dataset Statistics")
data_table(doc, [
    ["Total sensor readings",  "104,370"],
    ["Sensor channels",        "7 (soil moisture, temperature, humidity, reservoir, N, P, K)"],
    ["Sampling interval",      "~2 seconds"],
    ["Database engine",        "PostgreSQL 15"],
    ["ORM",                    "Prisma 5 (TypeScript)"],
    ["Total ML decisions",     "104,370 (one per reading)"],
], ["Parameter", "Value"], hdr_bg=C_BLUE_DARK)

sub_heading(doc, "4.2  Descriptive Statistics  (3,401 sampled rows from Python analysis)")
data_table(doc, [
    ["Soil Moisture (%)",     "63.24", "9.88",  "41.35", "79.18"],
    ["Temperature (degC)",    "27.51", "0.56",  "26.30", "29.60"],
    ["Humidity (% RH)",       "63.71", "3.84",  "56.00", "80.00"],
    ["Reservoir Level (%)",   "79.02", "25.81", "0.94",  "100.00"],
    ["Nitrogen (mg/kg)",      "18.28", "7.26",  "1.00",  "55.00"],
    ["Phosphorus (mg/kg)",    "24.74", "10.24", "1.00",  "55.00"],
    ["Potassium (mg/kg)",     "47.37", "22.40", "3.00",  "108.00"],
], ["Sensor", "Mean", "Std Dev", "Min", "Max"], hdr_bg=C_BLUE_DARK)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
#  5. DASHBOARD
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "5.", "Dashboard — Main Page  (/dashboard)")
add_para(doc,
    "The AquaSense dashboard is a Next.js 14 server component rendered at /dashboard. "
    "It auto-refreshes every 5 seconds. All DB fetches run in parallel via Promise.all() "
    "to minimise load time.", size=11, space_after=8)

sub_heading(doc, "5.1  Stats Bar — Four KPI Cards")
add_bullet(doc, "Total DB Records — live count from PostgreSQL (e.g. 104,370)")
add_bullet(doc, "Nodes Online — connected nodes vs total (e.g. 1 / 1)")
add_bullet(doc, "Active Alerts — unacknowledged threshold breaches (red when > 0)")
add_bullet(doc, "Sensors Tracking — fixed 8 channels listed below the value")

screenshot_box(doc,
    "Dashboard — Hero Header + Four KPI Stat Cards",
    "Navigate to localhost:3000/dashboard  ->  screenshot the top section")

sub_heading(doc, "5.2  Live Sensor Overview — System Averages")
add_para(doc,
    "Seven sensor cards in a row show system-wide averages across all online nodes. "
    "Each card has a coloured icon, current value, unit, and a status badge "
    "(Optimal / Low / High / Critical).", size=11, space_after=6)

screenshot_box(doc,
    "Dashboard — Live Sensor Overview Row (7 sensor cards)",
    "Screenshot of the 7 coloured average sensor cards")

sub_heading(doc, "5.3  Multi-Sensor Time-Series Chart")
add_para(doc,
    "A Recharts LineChart plots all 7 channels on a normalised 0-100 axis "
    "covering the last 24 hours of real DB data (or 720 mock points as fallback).", size=11, space_after=6)

screenshot_box(doc,
    "Dashboard — Multi-Sensor 24-Hour Time-Series Chart",
    "Screenshot of the 'Sensor Readings Over Time' chart section")

sub_heading(doc, "5.4  Conversion Formulas Panel")
add_para(doc,
    "Expandable panel displaying every ADC-to-unit conversion formula. "
    "Satisfies the assignment requirement to document signal conversion visibly.", size=11, space_after=6)

screenshot_box(doc,
    "Dashboard — Conversion Formulas Panel",
    "Screenshot of 'How Raw Sensor Signals Become Real Numbers'")

sub_heading(doc, "5.5  Data Classification & Analysis Panel")
add_para(doc,
    "Each reading is classified into criticalLow / low / optimal / high / criticalHigh. "
    "The panel shows per-sensor class distribution and a Pearson correlation matrix.", size=11, space_after=6)

screenshot_box(doc,
    "Dashboard — Data Classification & Analysis Panel",
    "Screenshot of the classification panel with distribution bars and correlation matrix")

sub_heading(doc, "5.6  AI Irrigation Intelligence Summary Card")
add_para(doc,
    "Compact card: latest ML decision, confidence %, pump state, model metrics "
    "(Accuracy, F1), total decisions, and a colour-coded distribution bar. "
    "Links to the full /ml page.", size=11, space_after=6)

screenshot_box(doc,
    "Dashboard — AI Irrigation Intelligence Summary Card",
    "Screenshot of the 'AI Irrigation Intelligence' section on /dashboard")

sub_heading(doc, "5.7  Individual Nodes + Pump Control")
add_para(doc,
    "NodeStatusGrid of clickable node cards alongside a PumpControlPanel with "
    "manual ON/OFF buttons and current source (ML / manual).", size=11, space_after=6)

screenshot_box(doc,
    "Dashboard — Node Status Grid + Pump Control Panel",
    "Screenshot of the 'Individual Sensor Nodes' section")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
#  6. NODE DETAIL
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "6.", "Individual Node Detail Page  (/nodes/[id])")
add_para(doc,
    "Clicking any node card opens a dedicated detail page showing 8 individual metric "
    "cards (2 rows of 4), a full AllSensorsChart with a time-range picker "
    "(1h / 6h / 24h / 7d), and a ClassificationPanel with calibration documentation.",
    size=11, space_after=8)

screenshot_box(doc,
    "Node Detail — Full Page View",
    "Navigate to localhost:3000/nodes/<id> and screenshot the full page")

screenshot_box(doc,
    "Node Detail — 8 Individual Metric Cards (2 rows of 4)",
    "Close-up of the 8 metric cards for one sensor node")

screenshot_box(doc,
    "Node Detail — AllSensors Chart + Time-Range Picker",
    "Screenshot of the multi-line chart with the 1h/6h/24h/7d picker visible")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
#  7. ML PAGE
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "7.", "AI Irrigation Intelligence Page  (/ml)")
add_para(doc,
    "The /ml page is the centrepiece of the machine-learning subsystem. "
    "It auto-refreshes every 6 seconds and displays the full output of the ML service, "
    "Python analysis pipeline, and decision history.", size=11, space_after=8)

sub_heading(doc, "7.1  Page Header and Stats Bar")
add_bullet(doc, "Model status badge — pulsing dot: Model Active / Model Offline")
add_bullet(doc, "Total Decisions, Model Accuracy, Irrigate Rate, Training Records KPI cards")

screenshot_box(doc,
    "ML Page — Hero Header + 4 KPI Stat Cards",
    "Navigate to localhost:3000/ml  ->  screenshot the top section")

sub_heading(doc, "7.2  Current Irrigation Decision (Live)")
add_para(doc,
    "MLHeroCard shows: decision label (colour-coded), confidence %, pump command, "
    "sensor inputs, timestamp, and per-class probabilities.", size=11, space_after=6)

screenshot_box(doc,
    "ML Page — Current Irrigation Decision Card (Live)",
    "Screenshot of the 'Current Irrigation Decision' card")

sub_heading(doc, "7.3  Live Arduino Command")
add_para(doc,
    "Displays the exact command string sent to the Arduino: "
    "CMD:PUMP=1:IRRIGATE or CMD:PUMP=0:HOLD / CMD:PUMP=0:LOW_WATER.", size=11, space_after=6)

screenshot_box(doc,
    "ML Page — Live Arduino Command Panel",
    "Screenshot of the 'Live Arduino Command' section")

sub_heading(doc, "7.4  Model Performance Metrics")
add_para(doc, "Evaluated on the 20% held-out test set:", size=11, space_after=4)
data_table(doc, [
    ["Accuracy",  "100.0 %", "Fraction of correct predictions"],
    ["Precision", "100.0 %", "Weighted average across 3 classes"],
    ["Recall",    "100.0 %", "Weighted average across 3 classes"],
    ["F1 Score",  "100.0 %", "Weighted average across 3 classes"],
], ["Metric", "Value", "Notes"], hdr_bg=C_GREEN_DARK)

info_box(doc,
    "Note:",
    "100% figures arise because the irrigation labels are derived from a deterministic "
    "two-feature rule (reservoir < 15% -> LOW_WATER; soil < 78% -> IRRIGATE; else HOLD). "
    "The cross-validation analysis (predicting soilMoisture_class from other sensors) "
    "gives a more challenging 94.99% mean F1.",
    bg=C_BLUE_LIGHT, border=C_BLUE_MID)

screenshot_box(doc,
    "ML Page — Model Performance Metrics Row",
    "Screenshot of the Accuracy, Precision, Recall, F1 metric tiles")

sub_heading(doc, "7.5  Feature Importance and Decision Distribution")
data_table(doc, [
    ["Reservoir Level",   "0.400  (40.0%)", "Primary driver — determines LOW_WATER class"],
    ["Soil Moisture",     "0.323  (32.3%)", "Secondary — IRRIGATE vs HOLD boundary"],
    ["Humidity",          "0.078  (7.8%)"],
    ["Potassium",         "0.073  (7.3%)"],
    ["Phosphorus",        "0.070  (7.0%)"],
    ["Nitrogen",          "0.039  (3.9%)"],
    ["Temperature",       "0.017  (1.7%)", "Least important for this task"],
], ["Feature", "Importance", "Interpretation"], hdr_bg=C_BLUE_DARK)

data_table(doc, [
    ["IRRIGATE",   "79,162", "75.9%", "Pump ON  — soil moisture below 78%"],
    ["LOW_WATER",  "12,996", "12.5%", "Pump OFF — reservoir below 15%"],
    ["HOLD",       "12,212", "11.7%", "Pump OFF — soil sufficiently wet"],
], ["Decision", "Count", "Share", "Meaning"], hdr_bg=C_GREEN_DARK)

screenshot_box(doc,
    "ML Page — Feature Importance + Decision Distribution panels",
    "Screenshot of both panels side-by-side on /ml")

sub_heading(doc, "7.6  Decision History Table")
add_para(doc,
    "Last 100 decisions: timestamp, node, decision badge, confidence, pump command, "
    "soil moisture, temperature, humidity, reservoir level, model version.", size=11, space_after=6)

screenshot_box(doc,
    "ML Page — Decision History Table (last 100 decisions)",
    "Screenshot of the 'Decision History' table section")

sub_heading(doc, "7.7  Full Python Analysis Charts")
add_para(doc,
    "After every retrain the ML service runs run_all.py, regenerating 13 PNG charts "
    "displayed in a responsive grid on /ml (click any to enlarge):", size=11, space_after=6)

for grp, charts in [
    ("ML Performance",
     ["07 — RF Confusion Matrix", "08 — Feature Importance Bar",
      "09 — GBR Regression Scatter", "11 — 5-model F1 Comparison",
      "12 — TensorFlow Training Curves", "13 — TensorFlow Confusion Matrix"]),
    ("Irrigation Decisions",
     ["10 — Decision Distribution Bar (PUMP_OFF / ON_LOW / ON_HIGH)"]),
    ("Exploratory Data Analysis",
     ["01 — Sensor Distributions", "02 — Class Distribution",
      "03 — Time-Series Overlay", "04 — Box-plots by Class",
      "05 — Soil vs Reservoir Scatter", "06 — Pearson Correlation Heatmap"]),
]:
    add_bullet(doc, "", bold_prefix=grp, color=C_BLUE_MID)
    for c in charts:
        add_bullet(doc, c, level=1)

screenshot_box(doc,
    "ML Page — Python Analysis Charts Grid (all 13 charts)",
    "Screenshot of the 'Full Analysis Report' chart grid on /ml")

sub_heading(doc, "7.8  How the AI Model Works")
add_para(doc,
    "MLHowItWorks explainer card covers: data ingestion, feature selection, "
    "label derivation, train/test split, Random Forest training, and inference flow.",
    size=11, space_after=6)

screenshot_box(doc,
    "ML Page — 'How the AI Model Works' Explainer Card",
    "Screenshot of the explainer section at the bottom of /ml")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
#  8. PYTHON PIPELINE
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "8.", "Python Analysis Pipeline")
add_para(doc,
    "An independent Python pipeline (python_analysis/) produces all charts, statistics, "
    "and CSV summaries. It is triggered automatically after every ML retrain and can also "
    "be run manually via python_analysis/run_all.py.", size=11, space_after=8)

sub_heading(doc, "8.1  Pipeline Modules")
data_table(doc, [
    ["00_export_data.py",      "Exports SensorReading from PostgreSQL -> data/raw_readings.csv"],
    ["01_cleaning.py",         "Deduplicates, clamps outliers, classifies readings -> cleaned.csv"],
    ["02_augmentation.py",     "Optional synthetic oversampling to balance classes"],
    ["03_summary_stats.py",    "mean/std/min/max/quartiles -> summary_statistics.csv"],
    ["04_eda.py",              "Plots 01-06: distributions, time-series, heatmap"],
    ["05_correlation.py",      "Pearson + Spearman matrices -> CSV files"],
    ["06_models.py",           "DT/RF/LR classification + Linear/GBR regression -> plots 07-09"],
    ["07_decision_logic.py",   "Rule-based irrigation logic -> decisions.csv + plot 10"],
    ["08_model_comparison.py", "5-model 5-fold cross-validation -> plot 11"],
    ["run_all.py",             "Runs all steps in order; --no-jupyter skips notebook export"],
], ["Module", "Purpose"], hdr_bg=C_BLUE_DARK)

sub_heading(doc, "8.2  Classification Model Results")
add_para(doc,
    "Predicting soilMoisture_class from 6 other sensor channels (80/20 stratified split):",
    size=11, space_after=4)
data_table(doc, [
    ["Random Forest",       "0.9397", "Best — selected for dashboard display"],
    ["Decision Tree",       "0.9155", "Good; slight overfit risk"],
    ["Logistic Regression", "0.7578", "Weakest — linear boundary insufficient"],
], ["Classifier", "Weighted F1 (test set)", "Notes"], hdr_bg=C_GREEN_DARK)

sub_heading(doc, "8.3  Regression Model Results")
add_para(doc, "Predicting raw soilMoisture percentage:", size=11, space_after=4)
data_table(doc, [
    ["Gradient Boosting", "5.019", "0.745", "Non-linear; preferred"],
    ["Linear Regression", "8.839", "0.210", "Weak — soil moisture is non-linear"],
], ["Regressor", "RMSE (%)", "R2", "Notes"], hdr_bg=C_BLUE_DARK)

sub_heading(doc, "8.4  5-Model Cross-Validation Comparison  (5-fold, stratified)")
data_table(doc, [
    ["Random Forest",       "0.9499", "0.0079", "0.9354", "0.9589", "Winner"],
    ["Gradient Boosting",   "0.9355", "0.0101", "0.9213", "0.9496", ""],
    ["Decision Tree",       "0.9302", "0.0171", "0.9054", "0.9505", ""],
    ["Naive Bayes",         "0.7712", "0.0126", "0.7576", "0.7936", ""],
    ["Logistic Regression", "0.7602", "0.0214", "0.7268", "0.7854", "Weakest"],
], ["Model", "Mean F1", "Std", "Min F1", "Max F1", "Rank"], hdr_bg=C_GREEN_DARK)

sub_heading(doc, "8.5  Generated Analysis Charts")
chart_list = [
    ("01", "Sensor Distributions",         "Histogram for each of the 7 channels"),
    ("02", "Class Distribution",           "Count per soilMoisture_class label"),
    ("03", "Time-Series Overlay",          "All 7 channels over full dataset time range"),
    ("04", "Box-plots by Class",           "Per-sensor spread across classification bands"),
    ("05", "Soil vs Reservoir Scatter",    "Primary irrigation features coloured by class"),
    ("06", "Pearson Correlation Heatmap",  "Pairwise correlations between all channels"),
    ("07", "RF Confusion Matrix",          "Random Forest classifier confusion matrix"),
    ("08", "RF Feature Importance",        "Feature importance bar chart from RF model"),
    ("09", "GBR Regression Scatter",       "Predicted vs actual soil moisture (GBR)"),
    ("10", "Decision Distribution",        "PUMP_OFF / PUMP_ON_LOW / PUMP_ON_HIGH counts"),
    ("11", "Model Comparison F1",          "5-model mean F1 with error bars"),
    ("12", "TensorFlow Training Curves",   "Loss and accuracy per epoch (Keras/TF)"),
    ("13", "TensorFlow Confusion Matrix",  "Confusion matrix for the TF classifier"),
]
for num, name, desc in chart_list:
    sub_sub_heading(doc, f"Chart {num} — {name}")
    add_para(doc, desc, size=10, color=C_GRAY_LIGHT, space_after=4)
    screenshot_box(doc,
        f"Python Analysis — {num} {name}",
        f"Insert from python_analysis/plots/{num}_*.png")

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
#  9. RESULTS
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "9.", "Results and Discussion")

sub_heading(doc, "9.1  Sensor Data Quality")
add_para(doc,
    "Over 104,370 readings were collected with all seven channels active. "
    "Soil moisture ranged from 41% (critically dry) to 79% (saturated), mean 63%. "
    "Reservoir level had high variance (std 25.8%) indicating genuine depletion/refill cycles. "
    "Temperature was tightly clustered (26.3-29.6 degC, std 0.56 degC) — consistent with "
    "a greenhouse environment.", size=11, space_after=8)

sub_heading(doc, "9.2  Machine Learning Performance")
add_para(doc,
    "The production Random Forest (ml_service) achieves 100% accuracy on its test set "
    "because labels follow a deterministic two-feature rule the model easily learns. "
    "The cross-validated analysis model (predicting soilMoisture_class from peer sensors) "
    "yields a more challenging 94.99% mean F1 — confirming genuine sensor relationships.", size=11, space_after=8)

sub_heading(doc, "9.3  Irrigation Decision Logic")
add_para(doc,
    "IRRIGATE dominates (75.9%) because the 78% soil threshold sits near the top of the "
    "observed range (max 79%). This was intentional: it makes ON/OFF transitions visible "
    "during demonstrations. LOW_WATER (12.5%) and HOLD (11.7%) occur at comparable rates.",
    size=11, space_after=8)

sub_heading(doc, "9.4  System Latency")
add_para(doc,
    "End-to-end latency (sensor -> pump command): ~2-4 s — Arduino 2 s loop + "
    "~100 ms HTTP POST + ~50 ms ML inference. Dashboard refreshes every 5-6 s, "
    "so every update appears within one cycle.", size=11, space_after=8)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
#  10. CONCLUSION
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "10.", "Conclusion")
add_para(doc, "AquaSense successfully demonstrates a complete IoT pipeline from physical sensor "
    "to AI-driven actuator control. Key achievements:", size=11, space_after=6)

conclusions = [
    "Seven sensors on one Arduino Uno with calibrated signal conversion formulas shown on the dashboard.",
    "104,370+ sensor readings stored in PostgreSQL — above the 100,000-record assignment target.",
    "Professional Next.js 14 dashboard with live auto-refresh, multi-sensor chart, alerts, and per-node pages.",
    "Random Forest ML service issuing real-time IRRIGATE/HOLD/LOW_WATER decisions and closing the loop via pump relay.",
    "Comprehensive Python analysis pipeline: 13 charts covering EDA, classification, regression, cross-validation, and TensorFlow.",
    "94.99% mean cross-validated F1 for the analysis model; 100% on the deterministic production task.",
    "Full feedback loop: Arduino -> Bridge -> Next.js -> PostgreSQL -> ML -> Dashboard -> Arduino.",
]
for c in conclusions:
    add_bullet(doc, c)

doc.add_paragraph().paragraph_format.space_after = Pt(8)
add_para(doc,
    "Future work: MQTT wireless communication, pH sensor integration, weather-forecast-aware "
    "scheduling, and cloud VM deployment for remote field monitoring.",
    italic=True, size=11, color=C_GRAY_LIGHT, space_after=8)

divider(doc)

# ════════════════════════════════════════════════════════════════════════════
#  11. REFERENCES
# ════════════════════════════════════════════════════════════════════════════

section_heading(doc, "11.", "References")
refs = [
    "Arduino LLC. (2024). Arduino Uno Reference Manual. arduino.cc/en/Main/ArduinoBoardUno",
    "Breiman, L. (2001). Random Forests. Machine Learning, 45(1), 5-32.",
    "Prisma Team. (2024). Prisma ORM Documentation. prisma.io/docs",
    "Vercel. (2024). Next.js 14 App Router Documentation. nextjs.org/docs",
    "FastAPI. (2024). FastAPI Framework Documentation. fastapi.tiangolo.com",
    "Scikit-learn Developers. (2024). sklearn.ensemble.RandomForestClassifier. scikit-learn.org",
    "Recharts. (2024). Recharts — Redefined chart library. recharts.org",
    "PostgreSQL Global Development Group. (2024). PostgreSQL 15 Documentation. postgresql.org/docs",
    "Aosong Electronics. (2021). DHT11 Humidity & Temperature Sensor Datasheet.",
    "ELECROW. (2022). HC-SR04 Ultrasonic Distance Sensor Datasheet.",
]
for i, ref in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.8)
    num_r = p.add_run(f"[{i}]  ")
    num_r.bold = True
    num_r.font.size = Pt(10)
    num_r.font.name = "Calibri"
    set_run_color(num_r, C_GREEN_DARK)
    txt_r = p.add_run(ref)
    txt_r.font.size = Pt(10)
    txt_r.font.name = "Calibri"
    set_run_color(txt_r, C_GRAY_MID)

# ════════════════════════════════════════════════════════════════════════════
#  SAVE
# ════════════════════════════════════════════════════════════════════════════

doc.save("AquaSense_Report.docx")
print("Report saved -> AquaSense_Report.docx")
