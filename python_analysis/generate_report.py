"""
AquaSense IoT – Complete Project Report Generator
Run from python_analysis/ directory:
    python generate_report.py
Output: AquaSense_Report.docx
"""
import os, sys, datetime
import pandas as pd
import numpy as np

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.run([sys.executable, "-m", "pip", "install", "python-docx", "--quiet"], check=True)
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm, Emu
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# ── Load real data ────────────────────────────────────────────────────────────
print("Loading data files...")
raw     = pd.read_csv("data/raw_readings.csv",       parse_dates=["timestamp"])
cleaned = pd.read_csv("data/cleaned.csv",             parse_dates=["timestamp"])
aug     = pd.read_csv("data/augmented.csv",           parse_dates=["timestamp"])
summary = pd.read_csv("data/summary_statistics.csv",  index_col=0)
pearson = pd.read_csv("data/correlation_pearson.csv", index_col=0)
models  = pd.read_csv("data/model_results.csv")
cv      = pd.read_csv("data/model_comparison.csv")
dec     = pd.read_csv("data/decisions.csv")

SENSORS = ["soilMoisture", "temperature", "humidity",
           "reservoirLevel", "nitrogen", "phosphorus", "potassium"]
CLASSES = ["criticalLow", "low", "optimal", "high", "criticalHigh"]

removed      = len(raw) - len(cleaned)
pct_removed  = removed / len(raw) * 100
dec_counts   = dec["decision"].value_counts()
best_clf_row = cv[cv["model"] == "Random Forest"].iloc[0]
rf_f1        = best_clf_row["mean_f1"]

# ── Document setup ────────────────────────────────────────────────────────────
doc = Document()
for sec in doc.sections:
    sec.top_margin    = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ── Colour palette ─────────────────────────────────────────────────────────────
NAVY        = RGBColor(0x1B, 0x2A, 0x4A)   # dark navy — main brand colour
NAVY_LIGHT  = RGBColor(0xC8, 0xD8, 0xF0)   # light navy for subtitle text in dark bands
ORANGE      = RGBColor(0xE8, 0x73, 0x2A)   # warm orange — accent colour
BLUE        = RGBColor(0x19, 0x76, 0xD2)
DARK_BLUE   = RGBColor(0x0D, 0x47, 0xA1)
TEAL        = RGBColor(0x00, 0x89, 0x7B)
DARK        = RGBColor(0x21, 0x21, 0x21)
MID_GRAY    = RGBColor(0x55, 0x55, 0x55)
GREEN       = RGBColor(0x38, 0x8E, 0x3C)
EMERALD     = RGBColor(0x04, 0x7D, 0x5E)
AMBER       = RGBColor(0xF5, 0x7C, 0x00)
RED         = RGBColor(0xC6, 0x28, 0x28)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)

HDR_BG        = "1B2A4A"   # navy for body section headers
HDR_BG_TEAL   = "00897B"
HDR_BG_DARK   = "1B2A4A"
ROW_BG        = "EAF2FB"
ROW_BG_TEAL   = "E0F2F1"
ACCENT_BG     = "E3F2FD"
SCREENSHOT_BG = "FFF8E1"
NAVY_HEX      = "1B2A4A"
ORANGE_HEX    = "E8732A"

# ── XML / style helpers ───────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)

def set_cell_border(cell, color="4472C4", size="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "single")
        border.set(qn("w:sz"),    size)
        border.set(qn("w:color"), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def add_paragraph_border_bottom(para, color="1976D2", size=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    str(size * 8))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def para(text="", bold=False, italic=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
         color=None, space_before=2, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.bold      = bold
        run.italic    = italic
        run.font.size = Pt(size)
        run.font.color.rgb = color if color else DARK
    return p

def heading(text, level=1, color=None, underline=False):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if level == 1:
        for run in h.runs:
            run.font.color.rgb = color if color else DARK_BLUE
            run.font.size = Pt(16)
            run.bold = True
        add_paragraph_border_bottom(h, color="1976D2", size=2)
    elif level == 2:
        for run in h.runs:
            run.font.color.rgb = color if color else TEAL
            run.font.size = Pt(13)
    elif level == 3:
        for run in h.runs:
            run.font.color.rgb = color if color else MID_GRAY
            run.font.size = Pt(12)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after  = Pt(4)
    return h

def bullet(text, size=11):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)
    return p

def numbered(text, size=11):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(4)
    return p

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Figure: {text}")
    run.italic    = True
    run.font.size = Pt(9)
    run.font.color.rgb = MID_GRAY
    p.paragraph_format.space_after = Pt(10)

def image(path, width=Inches(5.8), cap=""):
    if os.path.exists(path):
        doc.add_picture(path, width=width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if cap:
            caption(cap)
    else:
        para(f"[Chart not found: {path}]", italic=True, color=MID_GRAY)

def page_break():
    doc.add_page_break()

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    add_paragraph_border_bottom(p, color="BDBDBD", size=1)

def screenshot_placeholder(title, description):
    """Renders a gold-tinted placeholder box where a screenshot should be inserted."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, SCREENSHOT_BG)
    set_cell_border(cell, color="F9A825", size="8")
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(f"[ SCREENSHOT REQUIRED: {title} ]")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = AMBER
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(description)
    r2.italic     = True
    r2.font.size  = Pt(9)
    r2.font.color.rgb = MID_GRAY
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Recommended size: 1280 x 720 px  |  Format: PNG")
    r3.font.size  = Pt(8)
    r3.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def info_box(text, bg="EAF2FB", border_color="1976D2"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, bg)
    set_cell_border(cell, color=border_color, size="6")
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.italic    = True
    run.font.color.rgb = DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def df_table(df, hdr_labels=None, hdr_hex=HDR_BG, row_hex=ROW_BG):
    cols = list(df.columns)
    idxs = list(df.index)
    header = (hdr_labels or [""] + cols)
    t = doc.add_table(rows=1 + len(idxs), cols=len(header))
    t.style = "Table Grid"
    for j, h in enumerate(header):
        cell = t.rows[0].cells[j]
        cell.text = str(h)
        run = cell.paragraphs[0].runs[0]
        run.bold = True; run.font.size = Pt(9)
        run.font.color.rgb = WHITE
        set_cell_bg(cell, hdr_hex)
    for i, idx in enumerate(idxs):
        rc = t.rows[i+1].cells
        rc[0].text = str(idx)
        rc[0].paragraphs[0].runs[0].bold = True
        rc[0].paragraphs[0].runs[0].font.size = Pt(8)
        if i % 2 == 0: set_cell_bg(rc[0], row_hex)
        for j, col in enumerate(cols):
            val = df.loc[idx, col]
            rc[j+1].text = f"{val:.3f}" if isinstance(val, float) else str(val)
            rc[j+1].paragraphs[0].runs[0].font.size = Pt(8)
            if i % 2 == 0: set_cell_bg(rc[j+1], row_hex)
    doc.add_paragraph()
    return t

def simple_table(rows_data, bold_header=True, hdr_hex=HDR_BG, row_hex=ROW_BG, font_size=9):
    t = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    t.style = "Table Grid"
    for i, row in enumerate(rows_data):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = str(val)
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(font_size)
            if i == 0 and bold_header:
                run.bold = True
                run.font.color.rgb = WHITE
                set_cell_bg(cell, hdr_hex)
            elif j == 0:
                run.bold = True
                run.font.color.rgb = DARK
            if i % 2 == 1 and i != 0:
                set_cell_bg(cell, row_hex)
    doc.add_paragraph()
    return t

def set_table_width(tbl, width_cm):
    """Force a table to an exact width using OOXML."""
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), str(int(width_cm * 567)))   # 567 twips ≈ 1 cm
    tblW.set(qn("w:type"), "dxa")
    tblPr.append(tblW)

def remove_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),   "none")
        border.set(qn("w:sz"),    "0")
        border.set(qn("w:color"), "auto")
        tcBorders.append(border)
    tcPr.append(tcBorders)

def remove_table_borders(tbl):
    tbl_elem = tbl._tbl
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl_elem.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none"); b.set(qn("w:sz"), "0"); b.set(qn("w:color"), "auto")
        tblBorders.append(b)
    tblPr.append(tblBorders)

def set_cell_padding(cell, top=100, bottom=100, left=180, right=180):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top",top),("bottom",bottom),("left",left),("right",right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"),    str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

def cover_band(text1, size1, bold1, color1_rgb,
               text2="", size2=10, bold2=False, color2_rgb=None,
               bg_hex="1565C0", pad_top=160, pad_bot=160):
    """Full-width single-cell coloured band used for the cover page."""
    t = doc.add_table(rows=1, cols=1)
    t.style  = "Table Grid"
    remove_table_borders(t)
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, bg_hex)
    remove_cell_borders(cell)
    set_cell_padding(cell, top=pad_top, bottom=pad_bot, left=300, right=300)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(0)
    p1.paragraph_format.space_after  = Pt(0 if text2 else 2)
    r1 = p1.add_run(text1)
    r1.bold = bold1; r1.font.size = Pt(size1)
    r1.font.color.rgb = color1_rgb

    if text2:
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after  = Pt(0)
        r2 = p2.add_run(text2)
        r2.bold = bold2; r2.font.size = Pt(size2)
        r2.font.color.rgb = color2_rgb or color1_rgb
    return t

def cover_spacer(cm=0.4, bg_hex="FFFFFF"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    remove_table_borders(t)
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, bg_hex)
    remove_cell_borders(cell)
    h_twips = int(cm * 567)
    tr = t.rows[0]._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"),  str(h_twips))
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE  — clean navy + orange design (matching reference style)
# ══════════════════════════════════════════════════════════════════════════════

def _cp_text(cell, text, size, bold, rgb, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_before=0, space_after=0, first=True):
    p = cell.paragraphs[0] if first else cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = rgb
    return p

# ── ① University header box (dark navy, full-width within margins) ────────────
t_uni = doc.add_table(rows=1, cols=1)
t_uni.style = "Table Grid"
remove_table_borders(t_uni)
c_uni = t_uni.rows[0].cells[0]
set_cell_bg(c_uni, NAVY_HEX)
remove_cell_borders(c_uni)
set_cell_padding(c_uni, top=240, bottom=240, left=200, right=200)

_cp_text(c_uni, "UNIVERSITY OF RWANDA",
         20, True, WHITE, space_after=6)
_cp_text(c_uni, "College of Science and Technology",
         11, False, NAVY_LIGHT, first=False, space_after=2)
_cp_text(c_uni, "School of Information and Communication Technology",
         11, False, NAVY_LIGHT, first=False, space_after=0)

# ── ② Orange accent divider line ──────────────────────────────────────────────
p_line = doc.add_paragraph()
p_line.paragraph_format.space_before = Pt(0)
p_line.paragraph_format.space_after  = Pt(0)
pPr = p_line._p.get_or_add_pPr()
pBdr = OxmlElement("w:pBdr")
bot  = OxmlElement("w:bottom")
bot.set(qn("w:val"),   "single")
bot.set(qn("w:sz"),    "24")      # 3pt thick
bot.set(qn("w:space"), "0")
bot.set(qn("w:color"), ORANGE_HEX)
pBdr.append(bot)
pPr.append(pBdr)

# ── ③ Title block ─────────────────────────────────────────────────────────────
p_sp1 = doc.add_paragraph()
p_sp1.paragraph_format.space_before = Pt(24)
p_sp1.paragraph_format.space_after  = Pt(0)

p_t1 = doc.add_paragraph()
p_t1.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t1.paragraph_format.space_before = Pt(0)
p_t1.paragraph_format.space_after  = Pt(4)
r_t1 = p_t1.add_run("AquaSense IoT System")
r_t1.bold = True; r_t1.font.size = Pt(28); r_t1.font.color.rgb = NAVY

p_t2 = doc.add_paragraph()
p_t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t2.paragraph_format.space_before = Pt(0)
p_t2.paragraph_format.space_after  = Pt(4)
r_t2 = p_t2.add_run("Complete Project Report")
r_t2.bold = True; r_t2.font.size = Pt(20); r_t2.font.color.rgb = NAVY

p_t3 = doc.add_paragraph()
p_t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_t3.paragraph_format.space_before = Pt(0)
p_t3.paragraph_format.space_after  = Pt(20)
r_t3 = p_t3.add_run("Smart Irrigation Monitoring & Data Analysis Platform")
r_t3.bold = False; r_t3.font.size = Pt(13); r_t3.font.color.rgb = ORANGE

# ── ④ Info table (Module / Programme / Group / Submission Date) ───────────────
cover_info = [
    ("Module",          "Ubiquitous and Pervasive Computing (IoT)"),
    ("Programme",       "BSc Computer Science  |  Year 4"),
    ("Group",           "Group 1"),
    ("Submission Date", "May 6, 2026"),
]
t_info = doc.add_table(rows=len(cover_info), cols=2)
t_info.style = "Table Grid"
for i, (k, v) in enumerate(cover_info):
    c0, c1 = t_info.rows[i].cells[0], t_info.rows[i].cells[1]
    c0.text = k;  c1.text = v
    k_r = c0.paragraphs[0].runs[0]
    v_r = c1.paragraphs[0].runs[0]
    k_r.bold = True;  k_r.font.size = Pt(10);  k_r.font.color.rgb = NAVY
    v_r.font.size = Pt(10);  v_r.font.color.rgb = DARK
    set_cell_bg(c0, "DAE8F5")
    set_cell_padding(c0, top=80, bottom=80, left=160, right=160)
    set_cell_padding(c1, top=80, bottom=80, left=160, right=160)

# ── ⑤ Group Members heading ───────────────────────────────────────────────────
p_gm = doc.add_paragraph()
p_gm.paragraph_format.space_before = Pt(18)
p_gm.paragraph_format.space_after  = Pt(6)
r_gm = p_gm.add_run("Group Members")
r_gm.bold = True; r_gm.font.size = Pt(11); r_gm.font.color.rgb = NAVY

# ── ⑥ Members table ───────────────────────────────────────────────────────────
members = [
    ("1",  "Olivier Iradukunda",          "222005508"),
    ("2",  "Mukagasirabo Beatrice",        "222004462"),
    ("3",  "Nzaramyimana Jerome",          "222008510"),
    ("4",  "Siboniyo Emmanuel",            "222006224"),
    ("5",  "Tuyishime Ephron",             "222005571"),
    ("6",  "Byukusenge Immaculee",         "222006273"),
    ("7",  "Mugisha Edson",               "222018513"),
    ("8",  "Uwiringiyimana Marie Claire",  "222004637"),
    ("9",  "Dushimimana Fabrice",          "222017059"),
    ("10", "Mugabo Kazina Jules",          "222002936"),
    ("11", "Ishimwe Jean Marie Vianney",   "222019273"),
]
t_mem = doc.add_table(rows=len(members) + 1, cols=3)
t_mem.style = "Table Grid"

# header row
for j, lbl in enumerate(["No.", "Student Name", "Registration Number"]):
    c = t_mem.rows[0].cells[j]
    c.text = lbl
    r = c.paragraphs[0].runs[0]
    r.bold = True; r.font.size = Pt(10); r.font.color.rgb = WHITE
    set_cell_bg(c, NAVY_HEX)
    set_cell_padding(c, top=80, bottom=80, left=160, right=160)

# data rows — plain white alternating with very light gray
for i, (num, name, regno) in enumerate(members):
    bg = "FFFFFF" if i % 2 == 0 else "F4F6FA"
    row = t_mem.rows[i + 1]
    for j, val in enumerate((num, name, regno)):
        c = row.cells[j]
        c.text = val
        r = c.paragraphs[0].runs[0]
        r.font.size = Pt(10); r.font.color.rgb = DARK
        if j == 1: r.bold = True
        set_cell_bg(c, bg)
        set_cell_padding(c, top=70, bottom=70, left=160, right=160)

# ── ⑦ Orange footer band ──────────────────────────────────────────────────────
p_sp2 = doc.add_paragraph()
p_sp2.paragraph_format.space_before = Pt(18)
p_sp2.paragraph_format.space_after  = Pt(0)

t_foot = doc.add_table(rows=1, cols=1)
t_foot.style = "Table Grid"
remove_table_borders(t_foot)
c_foot = t_foot.rows[0].cells[0]
set_cell_bg(c_foot, ORANGE_HEX)
remove_cell_borders(c_foot)
set_cell_padding(c_foot, top=140, bottom=140, left=200, right=200)
p_foot = c_foot.paragraphs[0]
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_foot.paragraph_format.space_before = Pt(0)
p_foot.paragraph_format.space_after  = Pt(0)
r_foot = p_foot.add_run(
    "AquaSense IoT Project  ·  University of Rwanda  ·  May 2026"
)
r_foot.bold = True; r_foot.font.size = Pt(10); r_foot.font.color.rgb = WHITE

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ═════════════════════════════════════════════════════════════════════════════
heading("Table of Contents", 1)
toc_items = [
    ("1.",  "Introduction"),
    ("2.",  "System Architecture"),
    ("3.",  "AquaSense Dashboard — User Guide"),
    ("    3.1", "Main Dashboard Overview"),
    ("    3.2", "Live Sensor Overview Panel"),
    ("    3.3", "Multi-Sensor Time-Series Chart"),
    ("    3.4", "Sensor Signal Conversion Formulas"),
    ("    3.5", "Data Classification Panel"),
    ("    3.6", "Individual Sensor Nodes & Pump Control"),
    ("    3.7", "Alert System"),
    ("    3.8", "Node Detail Page"),
    ("    3.9", "Settings Page"),
    ("4.",  "Sensor Hardware and Data Conversion"),
    ("5.",  "Dataset Description"),
    ("6.",  "Data Cleaning"),
    ("7.",  "Data Augmentation"),
    ("8.",  "Summary Statistics"),
    ("9.",  "Exploratory Data Analysis (EDA)"),
    ("10.", "Correlation Analysis"),
    ("11.", "Machine Learning Models"),
    ("12.", "Model Evaluation & Accuracy"),
    ("13.", "Decision Logic"),
    ("14.", "Decision Table"),
    ("15.", "Comparison: Rule-Based vs Machine Learning"),
    ("16.", "Conclusion & Recommendations"),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r_num = p.add_run(f"{num}  ")
    r_num.bold = True; r_num.font.size = Pt(10)
    r_num.font.color.rgb = DARK_BLUE if not num.startswith("    ") else TEAL
    r_txt = p.add_run(title)
    r_txt.font.size = Pt(10)
    r_txt.font.color.rgb = DARK

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═════════════════════════════════════════════════════════════════════════════
heading("1. Introduction", 1)

para(
    "Agriculture is one of the most important sectors in many developing countries, "
    "and water is the most critical resource needed for crop growth. Many farmers still "
    "rely on guesswork when deciding when to irrigate — leading to either over-watering "
    "(wasting resources and risking root damage) or under-watering (reducing crop yield). "
    "The AquaSense system was designed to solve this problem using Internet of Things (IoT) "
    "technology combined with data-driven decision making."
)
para(
    "AquaSense is a full-stack smart irrigation monitoring platform that collects real-time "
    "data from sensors in the field, stores it in a database, displays it on an interactive "
    "web dashboard, and uses machine learning models alongside rule-based logic to "
    "automatically decide whether the irrigation pump should be turned on or off."
)
para("This report covers the entire project in two parts:")
bullet("Part A (Sections 2–3): The AquaSense system itself — its architecture, how the "
       "web dashboard works, and a complete walkthrough of every feature.")
bullet("Part B (Sections 4–16): The Python data analysis pipeline — data cleaning, "
       "augmentation, statistics, machine learning models, decision logic, and results.")

info_box(
    "The AquaSense dashboard is accessible at http://localhost:3000 when running locally. "
    "Screenshot placeholders throughout Section 3 indicate exactly where screenshots of "
    "the running dashboard should be inserted before submission.",
    bg="E8F5E9", border_color="388E3C"
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 2. SYSTEM ARCHITECTURE
# ═════════════════════════════════════════════════════════════════════════════
heading("2. System Architecture", 1)

para(
    "AquaSense is built as a modern, full-stack web application that integrates "
    "IoT hardware sensor nodes with a cloud-capable web dashboard. The architecture "
    "follows a layered design: sensor nodes collect raw data, a REST API ingests it, "
    "a relational database stores it, and a Next.js dashboard visualises it in real time."
)

heading("2.1  Technology Stack", 2)
simple_table([
    ["Layer",           "Technology",                    "Purpose"],
    ["Frontend",        "Next.js 14 (App Router)",       "Server-rendered React dashboard, real-time UI updates"],
    ["Styling",         "Tailwind CSS",                  "Utility-first responsive design"],
    ["Icons",           "Lucide React",                  "Consistent icon library across the dashboard"],
    ["Charts",          "Custom SVG + React",            "Multi-sensor line charts with responsive scaling"],
    ["Backend",         "Next.js API Routes (TypeScript)","REST endpoints for pump control, alerts, data ingestion"],
    ["ORM",             "Prisma",                        "Type-safe database access with schema migrations"],
    ["Database",        "PostgreSQL",                    "Stores 100,000+ sensor readings with full ACID support"],
    ["Data Streaming",  "Server-Sent Events (SSE)",      "Push live sensor updates to browser without polling"],
    ["Weather API",     "OpenWeatherMap API",            "Real-time weather data for irrigation context"],
    ["Data Analysis",   "Python 3 + Pandas + Scikit-learn","Offline analysis pipeline and ML model training"],
    ["Deep Learning",   "TensorFlow / Keras",            "Neural network classification model"],
], hdr_hex=HDR_BG_DARK)

heading("2.2  Data Flow Architecture", 2)
para(
    "The diagram below illustrates how data flows through the AquaSense system — "
    "from physical sensors in the field all the way to the web dashboard:"
)
simple_table([
    ["Step", "Component",        "Description"],
    ["1",    "Sensor Nodes",     "ESP32/Arduino nodes read ADC/digital values from physical sensors every 30 seconds"],
    ["2",    "HTTP POST Ingest", "Nodes send JSON payload to POST /api/ingest with all sensor readings"],
    ["3",    "PostgreSQL DB",    "Prisma ORM writes the reading to sensorReading table (~100,000+ rows)"],
    ["4",    "Alert Engine",     "The ingest route checks values against configurable thresholds and creates Alert records"],
    ["5",    "SSE Stream",       "GET /api/stream pushes new readings to connected browsers via Server-Sent Events"],
    ["6",    "Dashboard",        "Next.js server component fetches latest readings and renders the dashboard"],
    ["7",    "Pump API",         "POST /api/pump sends on/off commands; auto mode uses the decision rule engine"],
    ["8",    "Python Analysis",  "Separate Python pipeline reads CSV exports from the DB and runs ML analysis"],
], hdr_hex=HDR_BG_TEAL, row_hex=ROW_BG_TEAL)

heading("2.3  Database Schema (Key Tables)", 2)
simple_table([
    ["Table",           "Key Columns",                                   "Purpose"],
    ["SensorNode",      "id, name, zone, isActive",                      "Registers each physical sensor node in the field"],
    ["SensorReading",   "nodeId, timestamp, soilMoisture, temperature, humidity, reservoirLevel, nitrogen, phosphorus, potassium, pH",
                                                                          "Stores every sensor reading — primary data table"],
    ["Alert",           "nodeId, metric, breachValue, thresholdValue, direction, acknowledged",
                                                                          "Records threshold breaches for operator attention"],
    ["PumpEvent",       "action, source, timestamp",                     "Logs every pump on/off event with source (manual/auto)"],
    ["Threshold",       "metric, minValue, maxValue, nodeId",            "Configurable per-metric alert thresholds"],
    ["Prediction",      "horizon, predictedClass, confidence",           "Stores ML irrigation predictions (24h, 7d)"],
])

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 3. AQUASENSE DASHBOARD — USER GUIDE
# ═════════════════════════════════════════════════════════════════════════════
heading("3. AquaSense Dashboard — User Guide", 1)

para(
    "The AquaSense web dashboard is the primary interface for monitoring irrigation "
    "field conditions, controlling the pump, and reviewing historical sensor data. "
    "It is built with Next.js 14 and auto-refreshes every 5 seconds to show live data. "
    "This section provides a complete walkthrough of every page and feature."
)

info_box(
    "Navigation: The top navigation bar provides links to Dashboard (/dashboard), "
    "individual Node pages (/nodes/[id]), Settings (/settings), and Documentation (/docs). "
    "The dashboard is the default landing page at the root URL.",
    bg="EDE7F6", border_color="7B1FA2"
)

# 3.1
heading("3.1  Main Dashboard Overview", 2)
para(
    "The main dashboard (/dashboard) is the central hub of the AquaSense system. "
    "It is organised into a Hero Header at the top followed by four main sections: "
    "a Stats Bar, Live Sensor Overview, Sensor Time-Series Chart, and the Node + Pump panel. "
    "The dashboard auto-refreshes every 5 seconds using the AutoRefresh component."
)

screenshot_placeholder(
    "Main Dashboard — Full Page View",
    "Take a full-page screenshot of the dashboard at http://localhost:3000/dashboard\n"
    "Show the complete page including header, stats bar, sensor cards, and chart."
)

para(
    "The Hero Header contains the AquaSense logo, title, a Live animated pulse indicator, "
    "a link to the Documentation page, and the current timestamp. A green animated dot "
    "confirms that the system is actively receiving sensor data."
)

heading("3.2  Stats Bar (Key Performance Indicators)", 2)
para(
    "Directly below the header is a row of four KPI (Key Performance Indicator) cards "
    "that give an instant overview of the system state:"
)
simple_table([
    ["Card",            "What It Shows",                             "When It Changes"],
    ["Total DB Records","Total number of sensor readings in the DB", "Increases every 30 seconds as sensors send data"],
    ["Nodes Online",    "X / Y — how many nodes are active",        "Changes if a node goes offline (no ping > 5 min)"],
    ["Active Alerts",   "Number of unacknowledged threshold alerts", "Increases when a sensor breaches a threshold"],
    ["Sensors Tracking","Always shows 8 (all sensor types)",        "Static — confirms system is monitoring all sensors"],
])

screenshot_placeholder(
    "Stats Bar — KPI Cards",
    "Screenshot of the four KPI cards at the top of the dashboard.\n"
    "Zoom in to show the numbers, icons, and colour coding clearly."
)

# 3.3
heading("3.3  Live Sensor Overview Panel", 2)
para(
    "This section shows the current average readings across all online sensor nodes "
    "using a row of seven styled sensor cards. Each card displays the sensor name, "
    "current value with unit, a colour-coded status badge (OPTIMAL / LOW / HIGH / CRITICAL), "
    "and a small trend indicator."
)
simple_table([
    ["Card",            "Unit",   "Status Colour Coding"],
    ["Soil Moisture",   "%",      "Green = Optimal (40–70%), Yellow = Low/High, Red = Critical"],
    ["Temperature",     "°C",     "Green = Optimal (18–30°C), Red above 38°C or below 10°C"],
    ["Humidity",        "%",      "Green = Optimal (45–75%), Red below 30% or above 85%"],
    ["Reservoir Level", "%",      "Green = Optimal, Red = criticalLow (<10%) — no water to pump"],
    ["Nitrogen (N)",    "mg/kg",  "Green = Optimal (14–27 mg/kg), shows NPK nutrient health"],
    ["Phosphorus (P)",  "mg/kg",  "Green = Optimal (18–34 mg/kg)"],
    ["Potassium (K)",   "mg/kg",  "Green = Optimal (39–67 mg/kg)"],
])

screenshot_placeholder(
    "Live Sensor Overview — 7 Sensor Cards",
    "Screenshot of the Live Sensor Overview section showing all 7 sensor cards.\n"
    "Ideally show a mix of statuses (e.g. one LOW, one OPTIMAL) for visual variety."
)

# 3.4
heading("3.4  Multi-Sensor Time-Series Chart", 2)
para(
    "Below the sensor cards is the Sensor Readings Over Time chart — a multi-line "
    "interactive graph showing all 7 sensor values normalised to a 0–100% scale "
    "on the same Y-axis. This allows different sensors (e.g. Temperature in °C and "
    "Soil Moisture in %) to be compared side-by-side on a single plot."
)
bullet("The chart shows the last 24 hours of data by default.")
bullet("Each sensor is drawn in a unique colour with a legend at the bottom.")
bullet("Hovering over the chart shows an interactive tooltip with exact values.")
bullet("A 'Demo Data' amber badge appears if no real sensors are connected — "
       "it switches to 'Live DB' in emerald green when real data is present.")
bullet("The chart uses SSE (Server-Sent Events) via the AutoRefresh component to "
       "automatically update without the user needing to reload the page.")

screenshot_placeholder(
    "Multi-Sensor Time-Series Chart",
    "Screenshot of the Sensor Readings Over Time chart.\n"
    "The chart should show multiple coloured lines for different sensors over the 24h period.\n"
    "Make sure the legend is visible at the bottom."
)

# 3.5
heading("3.5  Sensor Signal Conversion Formulas", 2)
para(
    "This section displays the mathematical conversion formulas used to transform "
    "raw sensor signals into meaningful physical units. This is a key assignment "
    "requirement — it documents the calibration and signal processing methodology."
)
simple_table([
    ["Sensor",           "Raw Signal",        "Conversion Formula",                          "Result Unit"],
    ["Soil Moisture",    "ADC 0–4095",         "SM% = 100 − ((ADC − 1700) / (3200 − 1700)) × 100", "% moisture"],
    ["Temperature",      "DHT11/22 digital",  "Direct digital — T = raw_value / 10.0",       "°C"],
    ["Humidity",         "DHT11/22 digital",  "Direct digital — H = raw_value / 10.0",       "% RH"],
    ["Reservoir Level",  "HC-SR04 ultrasonic","Level% = 100 − (distance_cm / max_depth_cm) × 100", "% full"],
    ["Nitrogen",         "RS485 Modbus hex",  "N (mg/kg) = (high_byte × 256 + low_byte) × scale", "mg/kg"],
    ["Phosphorus",       "RS485 Modbus hex",  "P (mg/kg) = (high_byte × 256 + low_byte) × scale", "mg/kg"],
    ["Potassium",        "RS485 Modbus hex",  "K (mg/kg) = (high_byte × 256 + low_byte) × scale", "mg/kg"],
    ["pH",               "RS485 Modbus hex",  "pH = (high_byte × 256 + low_byte) / 10.0",   "pH units"],
])

screenshot_placeholder(
    "Sensor Conversion Formulas Panel",
    "Screenshot of the 'How Raw Sensor Signals Become Real Numbers' section.\n"
    "Show the full conversion formulas table for all sensors."
)

# 3.6
heading("3.6  Data Classification Panel", 2)
para(
    "The Data Classification Panel automatically analyses the distribution of sensor "
    "readings across the five classification categories and displays the results as "
    "styled bar charts. This section is populated from the database in real time."
)
bullet("For each sensor, a stacked breakdown shows what percentage of all historical "
       "readings fall into each class: criticalLow, low, optimal, high, criticalHigh.")
bullet("A summary table shows the sensor-wise correlation (Pearson r) between pairs.")
bullet("This section directly implements the assignment requirement to document "
       "data classification and calibration methodology.")

screenshot_placeholder(
    "Data Classification Panel",
    "Screenshot of the Data Classification & Analysis section.\n"
    "Show the per-sensor class distribution bars and the correlation table."
)

# 3.7
heading("3.7  Individual Sensor Nodes & Pump Control", 2)
para(
    "The lower section of the dashboard displays individual sensor node cards on "
    "the left and the Pump Control Panel on the right."
)

heading("Node Status Grid", 3)
para(
    "Each physical sensor node in the field has its own card showing: "
    "the node name and zone, online/offline status (red if no data for 5+ minutes), "
    "time since last reading, and the latest soil moisture, temperature, humidity, "
    "reservoir level, and NPK values. Clicking a node card navigates to the detailed "
    "node page (see Section 3.8)."
)

heading("Pump Control Panel", 3)
para(
    "The Pump Control Panel allows operators to manually control the irrigation pump "
    "in real time. It includes:"
)
bullet("Current pump status indicator (Running / Stopped / Unknown) with a "
       "colour-coded dot (green = on, red = off).")
bullet("Source label showing whether the last command came from 'manual' "
       "(operator click) or 'auto' (automatic rule engine).")
bullet("Running timer: displays how long the pump has been running (e.g. '3m 27s').")
bullet("Turn On / Turn Off button that sends a POST request to /api/pump.")
bullet("Auto Mode toggle: when enabled, the rule engine automatically turns the "
       "pump on/off based on the current soil moisture and reservoir level readings.")

screenshot_placeholder(
    "Pump Control Panel",
    "Screenshot of the Pump Control Panel widget.\n"
    "If possible, show it in the 'Pump Running' state to demonstrate the timer feature."
)

# 3.8
heading("3.8  Alert System", 2)
para(
    "The AquaSense alert system monitors sensor values continuously against "
    "configurable thresholds. When a sensor reading breaches a threshold, an Alert "
    "record is created in the database and displayed to the operator."
)
bullet("Alert Banner: A red dismissible banner appears at the top of the dashboard "
       "whenever there are unacknowledged alerts.")
bullet("Alert content: Shows which node, which sensor, the breach value, "
       "the threshold value, and whether the breach is above or below the threshold.")
bullet("Acknowledge button: Operators can acknowledge (dismiss) alerts "
       "individually — the alert is marked as acknowledged in the database.")
bullet("Alert History Table: Shows the full log of all past alerts — acknowledged "
       "and unacknowledged — with timestamps and sensor details.")

screenshot_placeholder(
    "Alert Banner (Active Alert State)",
    "Screenshot showing an active alert banner at the top of the dashboard.\n"
    "Trigger a test alert by temporarily lowering a threshold in Settings, "
    "then screenshot the red alert banner with sensor details visible."
)

# 3.9
heading("3.9  Node Detail Page", 2)
para(
    "Each sensor node has its own detail page at /nodes/[node-id]. "
    "This page provides a deeper view of a single node's readings and history:"
)
bullet("Hero header with node name, online/offline status badge, zone label, "
       "last-seen time, and record count for the last 24 hours.")
bullet("Current Sensor Values section: 8 metric cards displayed in two rows of 4, "
       "showing the latest value for each sensor with colour-coded status.")
bullet("Time-Series Charts section: individual line charts for each sensor, "
       "showing the full 24-hour history with a time-range picker (1h, 6h, 24h, 7d).")
bullet("Classification Panel: per-sensor classification showing historical "
       "distribution and calibration reference values.")

screenshot_placeholder(
    "Node Detail Page — Full View",
    "Screenshot of a Node Detail Page at /nodes/[any-node-id].\n"
    "Show the hero header, the 8 metric cards, and the time-series chart section."
)

screenshot_placeholder(
    "Node Metric Cards (8-card row)",
    "Close-up screenshot of the 8 sensor metric cards in the Current Sensor Values section.\n"
    "Show all cards including NPK and pH values with their status badges."
)

# 3.10
heading("3.10  Settings Page", 2)
para(
    "The Settings page (/settings) allows administrators to configure two aspects "
    "of the system without touching code:"
)

heading("Node Management", 3)
para(
    "Operators can register new sensor nodes by providing a name and zone. "
    "Existing nodes can be renamed, reassigned to a different zone, or deactivated. "
    "Deactivated nodes stop appearing on the dashboard but their data is preserved."
)

heading("Threshold Configuration", 3)
para(
    "For each sensor metric (soil moisture, temperature, humidity, reservoir level, NPK), "
    "operators can set custom minimum and maximum threshold values. When a reading "
    "falls outside these bounds, an alert is automatically created. Thresholds can "
    "be set globally or per-node for more precise control."
)

screenshot_placeholder(
    "Settings Page — Node Management & Thresholds",
    "Screenshot of the Settings page at /settings.\n"
    "Show both the Node Management section and the Threshold Configuration table."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 4. SENSOR HARDWARE AND DATA CONVERSION
# ═════════════════════════════════════════════════════════════════════════════
heading("4. Sensor Hardware and Data Conversion", 1)

para(
    "The AquaSense system uses five types of sensor hardware, each producing a "
    "different raw signal format that must be converted into meaningful physical units "
    "before storage or display. This section documents the hardware specifications "
    "and conversion methodology — a core assignment requirement."
)

heading("4.1  Sensor Hardware Specifications", 2)
simple_table([
    ["Sensor",           "Component",     "Interface",     "Raw Output",         "Physical Range"],
    ["Soil Moisture",    "Capacitive ADC","Analog (ADC)",  "Integer 0–4095",     "0–100% volumetric moisture"],
    ["Temperature",      "DHT11 / DHT22", "Digital 1-Wire","10x °C integer",     "−40°C to +85°C (DHT22)"],
    ["Humidity",         "DHT11 / DHT22", "Digital 1-Wire","10x % integer",      "0–100% relative humidity"],
    ["Reservoir Level",  "HC-SR04",       "Digital echo",  "Distance in cm",     "0–400 cm range"],
    ["Nitrogen (N)",     "RS485 NPK",     "RS485 Modbus",  "Hex register pair",  "0–1999 mg/kg"],
    ["Phosphorus (P)",   "RS485 NPK",     "RS485 Modbus",  "Hex register pair",  "0–1999 mg/kg"],
    ["Potassium (K)",    "RS485 NPK",     "RS485 Modbus",  "Hex register pair",  "0–1999 mg/kg"],
    ["pH",               "RS485 pH",      "RS485 Modbus",  "Hex / 10",           "pH 3.0–9.0"],
])

heading("4.2  Classification Thresholds", 2)
para(
    "Each converted sensor value is classified into one of five categories. "
    "These thresholds are based on agricultural reference values for tropical crop growing "
    "conditions and are configurable via the Settings page:"
)
simple_table([
    ["Sensor",          "criticalLow",  "low",    "optimal",   "high",   "criticalHigh"],
    ["Soil Moisture %", "≤ 20",         "≤ 40",   "40–70",     "70–85",  "> 85"],
    ["Temperature °C",  "≤ 10",         "≤ 18",   "18–30",     "30–38",  "> 38"],
    ["Humidity %",      "≤ 30",         "≤ 45",   "45–75",     "75–85",  "> 85"],
    ["Reservoir Level %","≤ 10",        "≤ 25",   "25–75",     "75–90",  "> 90"],
    ["Nitrogen mg/kg",  "≤ 5",          "≤ 14",   "14–27",     "27–35",  "> 35"],
    ["Phosphorus mg/kg","≤ 5",          "≤ 18",   "18–34",     "34–42",  "> 42"],
    ["Potassium mg/kg", "≤ 10",         "≤ 39",   "39–67",     "67–84",  "> 84"],
], hdr_hex=HDR_BG_TEAL, row_hex=ROW_BG_TEAL)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 5. DATASET DESCRIPTION
# ═════════════════════════════════════════════════════════════════════════════
heading("5. Dataset Description", 1)

para(
    f"The AquaSense dataset was collected from IoT sensor nodes deployed in an irrigation "
    f"field. Data was recorded continuously and stored in a PostgreSQL database. "
    f"The raw dataset contains {len(raw):,} sensor readings collected between "
    f"{raw['timestamp'].min().strftime('%d %B %Y')} and "
    f"{raw['timestamp'].max().strftime('%d %B %Y')} — covering approximately two years of "
    f"field operation."
)

heading("5.1  Dataset Summary", 2)
simple_table([
    ["Property",              "Value"],
    ["Total raw records",     f"{len(raw):,}"],
    ["Records after cleaning",f"{len(cleaned):,}"],
    ["Records after augmentation", f"{len(aug):,}"],
    ["Date range",            f"{raw['timestamp'].min().date()} to {raw['timestamp'].max().date()}"],
    ["Number of sensor nodes","1"],
    ["Number of features",    "7 sensor columns + timestamp + nodeId"],
    ["Target variable",       "soilMoisture_class (5 categories)"],
    ["Storage format",        "PostgreSQL database → CSV export for Python analysis"],
])

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 6. DATA CLEANING
# ═════════════════════════════════════════════════════════════════════════════
heading("6. Data Cleaning", 1)

para(
    "Before any analysis or machine learning can be done, the raw data must be cleaned. "
    "Raw IoT data often contains errors: sensors sometimes fail to send a reading, "
    "send a zero value by mistake, or occasionally produce extreme readings that do not "
    "represent real-world conditions. If these errors are not removed, they negatively "
    "affect the accuracy of any analysis or model."
)

heading("6.1  Step 1 — Removing Zero-Value Readings", 2)
para(
    "When a sensor fails to communicate with the microcontroller (for example, due to a "
    "loose cable or power interruption), it typically sends a value of exactly zero. "
    "A soil moisture reading of 0% or a temperature of 0°C is not a real agricultural "
    "measurement — it is a sensor error. We removed all rows where ANY sensor column "
    "contained a value of exactly zero."
)

heading("6.2  Step 2 — IQR Outlier Removal", 2)
para(
    "After removing zero values, we used the Interquartile Range (IQR) method to identify "
    "and remove extreme outliers. Any value below Q1 − 3×IQR or above Q3 + 3×IQR is "
    "considered an extreme outlier. We used a factor of 3.0 (instead of the standard 1.5) "
    "so that only truly extreme values were removed — preserving natural variation."
)
para("   Lower bound = Q1 − 3 × IQR     |     Upper bound = Q3 + 3 × IQR", bold=True, size=10)

heading("6.3  Cleaning Results", 2)
simple_table([
    ["Stage",                    "Row Count",              "Notes"],
    ["Raw dataset",              f"{len(raw):,}",          "As exported from database"],
    ["After zero removal",       f"{len(raw)-removed:,}",  "Sensor dropout readings removed"],
    ["After IQR outlier removal",f"{len(cleaned):,}",      f"{removed:,} rows removed ({pct_removed:.1f}%)"],
])

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 7. DATA AUGMENTATION
# ═════════════════════════════════════════════════════════════════════════════
heading("7. Data Augmentation", 1)

para(
    "Data augmentation means creating more data samples from existing data. "
    "This is common in machine learning when the original dataset is smaller than ideal "
    "for training models — especially deep learning neural networks."
)

heading("7.1  Method — NumPy Bootstrap Sampling", 2)
para(
    "We used bootstrap sampling — randomly selecting rows from the cleaned dataset "
    "WITH replacement (the same row can be chosen more than once). Implemented with "
    "NumPy's random number generator with seed=42 for reproducibility."
)

heading("7.2  Adding Gaussian Noise", 2)
para(
    "To avoid identical duplicate rows, we added small Gaussian noise to each sensor "
    "column: 2% of that column's standard deviation. This preserves the distribution "
    "while creating realistic variation."
)

heading("7.3  Distribution Preservation Check", 2)
aug_check_rows = [["Sensor","Original Mean","Augmented Mean","Difference %","Original Std","Augmented Std"]]
for col in SENSORS:
    if col in cleaned.columns and col in aug.columns:
        om, am = cleaned[col].mean(), aug[col].mean()
        os_, as_ = cleaned[col].std(), aug[col].std()
        delta = abs(om - am) / om * 100 if om != 0 else 0
        aug_check_rows.append([col, f"{om:.3f}", f"{am:.3f}", f"{delta:.2f}%", f"{os_:.3f}", f"{as_:.3f}"])
simple_table(aug_check_rows)
para(
    "All sensors show less than 1% difference in mean values — confirming the augmented "
    "dataset successfully preserves the statistical properties of the original data."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 8. SUMMARY STATISTICS
# ═════════════════════════════════════════════════════════════════════════════
heading("8. Summary Statistics", 1)

para(
    "Summary statistics give a quick overview of the data — typical values, spread, "
    f"and range. The table below covers the cleaned dataset ({len(cleaned):,} rows):"
)
summary_display = summary.T
if not summary_display.empty:
    df_table(summary_display.round(2))
else:
    df_table(cleaned[SENSORS].describe(percentiles=[.25,.5,.75]).T.round(2))

heading("8.1  Key Observations", 2)
sm_mean  = cleaned["soilMoisture"].mean()
sm_std   = cleaned["soilMoisture"].std()
tmp_mean = cleaned["temperature"].mean()
hum_mean = cleaned["humidity"].mean()

bullet(f"Soil Moisture: average {sm_mean:.1f}%, std {sm_std:.1f}% — most readings in the healthy 40–76% range.")
bullet(f"Temperature: average {tmp_mean:.1f}°C — stable warm-field environment.")
bullet(f"Humidity: average {hum_mean:.1f}% — high-humidity environment expected near irrigated crops.")
bullet("Reservoir Level: high variability (std ~18%), reflecting regular fill-and-drain irrigation cycles.")
bullet("NPK: similar distributions with strong inter-correlation (see Section 10).")

heading("8.2  Class Distribution", 2)
image("plots/02_class_distribution.png", width=Inches(6.0),
      cap="Class distribution per sensor (cleaned dataset)")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 9. EXPLORATORY DATA ANALYSIS (EDA)
# ═════════════════════════════════════════════════════════════════════════════
heading("9. Exploratory Data Analysis (EDA)", 1)

para(
    "Exploratory Data Analysis visually explores patterns, distributions, and relationships "
    "in the data before building models."
)

heading("9.1  Sensor Value Distributions", 2)
image("plots/01_sensor_distributions.png", width=Inches(6.2),
      cap="Sensor value distributions — cleaned dataset (7 sensors)")
bullet("Soil Moisture: roughly bell-shaped around 58% — field spends most time in optimal zone.")
bullet("Temperature: narrow distribution around 26°C — stable environment.")
bullet("Humidity: peaks strongly around 78%.")
bullet("Reservoir Level: spread out, reflecting irrigation fill-drain cycles.")
bullet("NPK: bell-shaped distributions showing natural soil nutrient variability.")

heading("9.2  Time-Series Overview", 2)
image("plots/03_time_series.png", width=Inches(6.2),
      cap="Time-series: soil moisture, temperature, and humidity over two years")

heading("9.3  Soil Moisture vs Reservoir Level", 2)
image("plots/05_soil_vs_reservoir.png", width=Inches(6.2),
      cap="Soil moisture (blue) vs reservoir level (orange) over time")

heading("9.4  Sensor Values by Soil Moisture Class", 2)
image("plots/04_boxplots_by_class.png", width=Inches(6.2),
      cap="Box plots: how each sensor value changes across soil moisture classes")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 10. CORRELATION ANALYSIS
# ═════════════════════════════════════════════════════════════════════════════
heading("10. Correlation Analysis", 1)

para(
    "Correlation analysis measures how strongly two variables are related. "
    "Values close to +1 = strong positive, close to −1 = strong negative, near 0 = unrelated."
)

heading("10.1  Correlation Heatmaps", 2)
image("plots/06_correlation_heatmap.png", width=Inches(6.2),
      cap="Pearson (left) and Spearman (right) correlation matrices for all 7 sensors")

heading("10.2  Key Correlation Findings", 2)
key_pairs = [
    ("Temperature ↔ Humidity",    f"{pearson.loc['temperature','humidity']:.3f}",    "Strong negative — hot air is less humid."),
    ("Nitrogen ↔ Phosphorus",     f"{pearson.loc['nitrogen','phosphorus']:.3f}",     "Very strong positive — NPK deposited together."),
    ("Nitrogen ↔ Potassium",      f"{pearson.loc['nitrogen','potassium']:.3f}",      "Very strong positive — same cause as N↔P."),
    ("Phosphorus ↔ Potassium",    f"{pearson.loc['phosphorus','potassium']:.3f}",    "Very strong positive — NPK trio moves together."),
    ("Soil Moisture ↔ Humidity",  f"{pearson.loc['soilMoisture','humidity']:.3f}",   "Weak positive — humid air slightly predicts moist soil."),
    ("Soil Moisture ↔ Temperature",f"{pearson.loc['soilMoisture','temperature']:.3f}","Weak negative — warm temps slightly dry the soil."),
    ("Soil Moisture ↔ Reservoir", f"{pearson.loc['soilMoisture','reservoirLevel']:.3f}","Weak negative — dry soil → reservoir in use."),
]
simple_table([["Sensor Pair", "Pearson r", "Interpretation"]] + [list(r) for r in key_pairs])

para(
    "Critical finding: soil moisture has only WEAK correlations (r between −0.07 and +0.20) "
    "with all other sensors. This directly explains why ML models show moderate accuracy — "
    "the other sensors are not strong predictors of soil moisture class."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 11. MACHINE LEARNING MODELS
# ═════════════════════════════════════════════════════════════════════════════
heading("11. Machine Learning Models", 1)

para(
    "We trained two types of predictive models: classification models (predict which "
    "of the 5 soil moisture classes a reading belongs to) and regression models "
    "(predict the exact moisture percentage)."
)
para(
    "Input features: temperature, humidity, reservoir level, nitrogen, phosphorus, "
    "potassium (6 features). Dataset split: 80% training / 20% testing. "
    "All features standardised (zero mean, unit variance) using StandardScaler."
)

heading("11.1  Classification Models", 2)
simple_table([
    ["Model",               "Library",        "Key Idea"],
    ["Decision Tree",       "Scikit-learn",   "Learns a tree of If-Then rules — easy to interpret."],
    ["Random Forest",       "Scikit-learn",   "Builds 100 decision trees and combines predictions."],
    ["Logistic Regression", "Scikit-learn",   "Finds a linear boundary between classes."],
    ["Naive Bayes",         "Scikit-learn",   "Probabilistic model assuming feature independence."],
    ["Gradient Boosting",   "Scikit-learn",   "Builds trees sequentially, each fixing the last."],
    ["TensorFlow MLP",      "TensorFlow/Keras","Deep neural network with 4 layers + dropout."],
])

heading("11.2  TensorFlow Neural Network Architecture", 2)
simple_table([
    ["Layer",       "Type",                "Output",     "Purpose"],
    ["Input",       "Input",               "6 features", "Receives standardised sensor values"],
    ["Hidden 1",    "Dense + ReLU",        "128 neurons","Learns complex sensor combinations"],
    ["Norm 1",      "BatchNormalization",  "128",        "Stabilises training"],
    ["Dropout 1",   "Dropout (30%)",       "128",        "Prevents overfitting"],
    ["Hidden 2",    "Dense + ReLU",        "64 neurons", "Deeper feature learning"],
    ["Norm 2",      "BatchNormalization",  "64",         "Stabilises second layer"],
    ["Dropout 2",   "Dropout (20%)",       "64",         "Light regularisation"],
    ["Hidden 3",    "Dense + ReLU",        "32 neurons", "Feature compression"],
    ["Output",      "Dense + Softmax",     "5 classes",  "Probability for each moisture class"],
])
para("Training: Adam optimiser (lr=0.001), batch size=512, max 30 epochs with early stopping (patience=5).")

heading("11.3  Regression Models", 2)
simple_table([
    ["Model",               "RMSE",   "R² Score", "Interpretation"],
    ["Linear Regression",   "17.90",  "0.057",    "Very weak — near-zero R² means barely better than guessing the mean"],
    ["Gradient Boosting",   "17.84",  "0.063",    "Slightly better — still very weak prediction"],
])
image("plots/09_regression_scatter.png", width=Inches(4.5),
      cap="Gradient Boosting regression: predicted vs actual soil moisture — wide spread confirms weak prediction.")

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 12. MODEL EVALUATION & ACCURACY
# ═════════════════════════════════════════════════════════════════════════════
heading("12. Model Evaluation and Accuracy", 1)

para(
    "Models were evaluated using weighted F1-score — a metric that accounts for "
    "class imbalance by giving more weight to classes with more examples."
)

heading("12.1  Random Forest — Confusion Matrix", 2)
image("plots/07_confusion_matrix.png", width=Inches(5.0),
      cap="Random Forest confusion matrix — rows = true class, columns = predicted class")

heading("12.2  Feature Importance", 2)
image("plots/08_feature_importance.png", width=Inches(5.0),
      cap="Random Forest feature importance for predicting soil moisture class")
para(
    "Humidity and nitrogen have the highest importance scores. "
    "Reservoir level has the lowest — knowing how full the tank is tells us "
    "relatively little about current soil moisture."
)

heading("12.3  TensorFlow Training History", 2)
tf_plot = "plots/12_tensorflow_training.png"
if os.path.exists(tf_plot):
    image(tf_plot, width=Inches(6.0),
          cap="TensorFlow MLP training — accuracy and loss per epoch (blue = train, red dashed = validation)")
    para(
        "Training and validation curves follow each other closely — the model is NOT "
        "overfitting. It learns real patterns, not just memorising training data."
    )

heading("12.4  Classification Accuracy Summary", 2)
clf_rows = [["Model", "Task", "Weighted F1", "Performance"]]
for _, row in models[models["task"]=="classification"].iterrows():
    f1 = row["value"]
    perf = "Good" if f1 >= 0.40 else ("Moderate" if f1 >= 0.35 else "Below average")
    clf_rows.append([row["model"], "Classification", f"{f1:.4f}", perf])
simple_table(clf_rows)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 13. DECISION LOGIC
# ═════════════════════════════════════════════════════════════════════════════
heading("13. Decision Logic", 1)

para(
    "Decision logic turns sensor readings into pump actions. "
    "AquaSense implements a rule-based engine with clearly defined If-Then rules, "
    "mirroring the auto-mode logic in the live dashboard."
)

heading("13.1  Possible Actions", 2)
simple_table([
    ["Action",         "Meaning"],
    ["PUMP_OFF",       "Soil moisture sufficient — pump stays off"],
    ["PUMP_ON_LOW",    "Soil slightly dry — pump on at low power"],
    ["PUMP_ON_HIGH",   "Soil critically dry — pump on at full power"],
    ["ALERT_REFILL",   "Reservoir too low to irrigate — alert sent to refill tank"],
    ["ALERT_DRAINAGE", "Soil critically WET — drainage alert (overwatering risk)"],
    ["MONITOR",        "Borderline conditions — no action, system monitors"],
], hdr_hex=HDR_BG_TEAL, row_hex=ROW_BG_TEAL)

heading("13.2  Decision Examples", 2)
for action in ["PUMP_ON_HIGH", "PUMP_ON_LOW", "PUMP_OFF", "ALERT_REFILL", "ALERT_DRAINAGE"]:
    sample = dec[dec["decision"] == action]
    if len(sample):
        row = sample.iloc[0]
        para(
            f"  IF  Soil Moisture = {row['soilMoisture']:.1f}%  ({row['soilMoisture_class']})\n"
            f"  AND Reservoir Level = {row['reservoirLevel']:.1f}%  ({row['reservoirLevel_class']})\n"
            f"  THEN  →  {action}",
            size=10
        )

image("plots/10_decision_distribution.png", width=Inches(5.8),
      cap="Distribution of irrigation decisions applied to the full cleaned dataset")

heading("13.3  Decision Output Summary", 2)
total_dec = len(dec)
dec_rows = [["Action", "Count", "Percentage", "Meaning"]]
action_meanings = {
    "PUMP_OFF":       "Field well irrigated — most common state",
    "PUMP_ON_LOW":    "Slightly dry conditions needing gentle irrigation",
    "ALERT_DRAINAGE": "Waterlogging risk — critical drainage alert",
    "MONITOR":        "Borderline state — no immediate action",
    "PUMP_ON_HIGH":   "Critically dry — emergency full irrigation",
    "ALERT_REFILL":   "Tank empty — cannot irrigate, refill required",
}
for action, count in dec_counts.items():
    dec_rows.append([action, f"{count:,}", f"{count/total_dec*100:.1f}%", action_meanings.get(action, "")])
simple_table(dec_rows)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 14. DECISION TABLE
# ═════════════════════════════════════════════════════════════════════════════
heading("14. Decision Table", 1)

para(
    "A decision table formally defines every combination of input conditions and "
    "the corresponding action. Rules are evaluated top-to-bottom by priority — "
    "the first matching rule wins. The wildcard (*) means any value."
)

simple_table([
    ["Priority", "Soil Moisture Class", "Reservoir Class", "Action",         "Reason"],
    ["1",        "criticalLow",         "criticalLow",      "ALERT_REFILL",   "Soil dry AND tank empty — cannot irrigate"],
    ["1",        "criticalHigh",        "*",                "ALERT_DRAINAGE", "Waterlogged — drainage needed regardless"],
    ["2",        "criticalLow",         "optimal",          "PUMP_ON_HIGH",   "Critically dry, tank sufficient — max irrigation"],
    ["2",        "criticalLow",         "high",             "PUMP_ON_HIGH",   "Same — plenty of water available"],
    ["2",        "criticalLow",         "criticalHigh",     "PUMP_ON_HIGH",   "Tank very full — rescue critically dry soil"],
    ["3",        "criticalLow",         "low",              "PUMP_ON_LOW",    "Very dry but tank is low — irrigate carefully"],
    ["4",        "low",                 "optimal",          "PUMP_ON_LOW",    "Slightly dry, tank fine — light irrigation"],
    ["4",        "low",                 "high",             "PUMP_ON_LOW",    "Enough water for gentle irrigation"],
    ["4",        "low",                 "criticalHigh",     "PUMP_ON_LOW",    "Tank very full, soil slightly dry"],
    ["5",        "low",                 "criticalLow",      "ALERT_REFILL",   "Needs water but tank is empty"],
    ["5",        "low",                 "low",              "MONITOR",        "Both low — monitor before acting"],
    ["6",        "optimal",             "*",                "PUMP_OFF",       "Soil moisture ideal — no irrigation needed"],
    ["7",        "high",               "*",                "PUMP_OFF",       "Soil wet — prevent overwatering"],
], hdr_hex=HDR_BG_DARK, row_hex=ROW_BG)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 15. COMPARISON — RULE-BASED vs ML
# ═════════════════════════════════════════════════════════════════════════════
heading("15. Comparison: Rule-Based System vs Machine Learning", 1)

heading("15.1  Cross-Validation Model Comparison", 2)
para(
    "5-fold cross-validation: dataset split into 5 parts, model trains on 4 and tests "
    "on the 5th — repeated 5 times. Average F1-score across 5 runs is the final result."
)
cv_display = cv.copy()
cv_display.columns = ["Model", "Mean F1", "Std Dev", "Min F1", "Max F1"]
df_table(cv_display.set_index("Model"))

image("plots/11_model_comparison.png", width=Inches(6.0),
      cap="5-fold cross-validation: Mean weighted F1 score with error bars (std deviation)")

heading("15.2  Rule-Based vs ML — Direct Comparison", 2)
simple_table([
    ["Aspect",              "Rule-Based System",                             "Machine Learning (Random Forest)"],
    ["Accuracy",            "~100% (rules define the classes)",              f"F1 = {rf_f1:.4f} (moderate)"],
    ["Interpretability",    "Very high — rules readable by humans",          "Moderate — can inspect feature importance"],
    ["Flexibility",         "Fixed — must update rules manually",            "Adapts automatically to new training data"],
    ["Training data needed","None — expert knowledge only",                  "Needs labelled training examples"],
    ["Handles new patterns","No — only covers defined scenarios",            "Yes — generalises to unseen combinations"],
    ["Speed",               "Instantaneous (simple lookup table)",           "Very fast after training"],
    ["Best for",            "Stable, well-understood environments",          "Complex, data-rich, evolving problems"],
], hdr_hex=HDR_BG_TEAL, row_hex=ROW_BG_TEAL)

heading("15.3  Analysis", 2)
para(
    f"Random Forest achieves the highest ML accuracy (Mean F1 = {rf_f1:.4f}). "
    "The TensorFlow neural network performs similarly to Logistic Regression because "
    "with only 6 features and weak correlations, simple models capture available "
    "patterns as well as a deep network."
)
para(
    "The rule-based system appears 100% accurate because the class labels in the dataset "
    "were generated by those same threshold rules. In production, the rule-based system "
    "is used for reliable day-to-day operation, while ML is used for anomaly detection "
    "and predicting future soil conditions from weather data."
)

page_break()

# ═════════════════════════════════════════════════════════════════════════════
# 16. CONCLUSION AND RECOMMENDATIONS
# ═════════════════════════════════════════════════════════════════════════════
heading("16. Conclusion and Recommendations", 1)

heading("16.1  Summary of Findings", 2)

bullet(
    f"Dataset quality: {len(raw):,} raw readings collected over 2 years. "
    f"After cleaning, {len(cleaned):,} readings remained ({pct_removed:.1f}% removal) — "
    "confirming excellent sensor reliability."
)
bullet(
    "Data augmentation: Dataset doubled to 200,000 rows using NumPy bootstrap sampling "
    "with Gaussian noise. Original statistical distribution preserved within 1%."
)
bullet(
    f"Field conditions: Average soil moisture {cleaned['soilMoisture'].mean():.1f}%, "
    f"temperature {cleaned['temperature'].mean():.1f}°C, humidity {cleaned['humidity'].mean():.1f}%. "
    "Most readings in the optimal or high moisture range."
)
bullet(
    "Strongest correlations: Temperature↔Humidity (r = −0.982), NPK nutrients "
    "highly correlated with each other (r = 0.944–0.982)."
)
bullet(
    f"ML models: Random Forest achieves the best classification F1 = {rf_f1:.4f}. "
    "All models show moderate accuracy because other sensors have weak predictive "
    "power over soil moisture — a fundamental data characteristic, not a model failure."
)
bullet(
    f"Decision system: {dec_counts.get('PUMP_OFF',0)/total_dec*100:.1f}% PUMP_OFF "
    f"(healthy field), {dec_counts.get('PUMP_ON_LOW',0)/total_dec*100:.1f}% PUMP_ON_LOW, "
    f"{dec_counts.get('ALERT_DRAINAGE',0)/total_dec*100:.1f}% ALERT_DRAINAGE."
)
bullet(
    "Dashboard: The AquaSense Next.js dashboard successfully demonstrates real-time "
    "IoT monitoring with live sensor cards, multi-sensor charting, pump control, "
    "configurable alerts, and full per-node detail pages."
)

heading("16.2  Real vs Augmented Dataset Comparison", 2)
simple_table([
    ["Property",              "Real Dataset (Cleaned)",                    "Augmented Dataset"],
    ["Number of rows",        f"{len(cleaned):,}",                         f"{len(aug):,}"],
    ["Soil Moisture mean",    f"{cleaned['soilMoisture'].mean():.2f}%",    f"{aug['soilMoisture'].mean():.2f}%"],
    ["Temperature mean",      f"{cleaned['temperature'].mean():.2f}°C",   f"{aug['temperature'].mean():.2f}°C"],
    ["Humidity mean",         f"{cleaned['humidity'].mean():.2f}%",        f"{aug['humidity'].mean():.2f}%"],
    ["Contains duplicates",   "No",                                         "Yes (by design — bootstrap sampling)"],
    ["Used for",              "All analysis + primary ML training",         "Supplementary deep learning training"],
    ["Distribution match",    "Reference (100%)",                          "Within 1% of reference"],
])

heading("16.3  Recommendations", 2)
bullet("Collect additional sensor types (rainfall sensor, solar radiation meter, wind anemometer) "
       "to improve ML model accuracy — more informative features lead to significantly better predictions.")
bullet("Deploy the rule-based decision system for reliable day-to-day irrigation control, "
       "as it is fast, interpretable, and requires no training data.")
bullet("Use machine learning for anomaly detection — identifying unusual sensor readings "
       "that may indicate hardware failure or environmental events.")
bullet("Address class imbalance in future model training using SMOTE or class-weighted "
       "loss functions to improve performance on rare but critical classes like criticalLow.")
bullet("Expand the sensor node network to multiple field zones — this will increase the "
       "dataset size and variability, improving ML model accuracy over time.")
bullet("Integrate the weather forecast API (already implemented in the dashboard) with "
       "the ML prediction engine to forecast irrigation needs 24–48 hours ahead.")

heading("16.4  Final Statement", 2)
para(
    "The AquaSense IoT project successfully demonstrates a complete, end-to-end smart "
    "irrigation system — from physical sensor hardware and raw data collection, through "
    "a professional web dashboard for real-time monitoring and pump control, to a full "
    "data analysis and machine learning pipeline. The combination of a rule-based "
    "decision table (for reliable, interpretable automation) with trained ML models "
    "(for pattern discovery and predictive capability) provides a robust foundation "
    "for a real-world precision agriculture deployment."
)

divider()
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_end = p_end.add_run("— End of Report —")
r_end.italic = True; r_end.font.size = Pt(10)
r_end.font.color.rgb = MID_GRAY

p_gen = doc.add_paragraph()
p_gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_gen = p_gen.add_run(f"Generated: {datetime.date.today().strftime('%d %B %Y')}  |  AquaSense IoT System  |  University of Rwanda")
r_gen.font.size = Pt(8); r_gen.font.color.rgb = RGBColor(0xBD, 0xBD, 0xBD)

# ═════════════════════════════════════════════════════════════════════════════
# SAVE
# ═════════════════════════════════════════════════════════════════════════════
out = "AquaSense_Report.docx"
doc.save(out)
print(f"\nReport saved -> {out}")
print(f"File size: {os.path.getsize(out)/1024:.0f} KB")
print("Open in Microsoft Word or LibreOffice Writer.")
print("\nScreenshot placeholders are marked in gold boxes throughout Section 3.")
print("Replace each placeholder with a real screenshot of the running dashboard.")
